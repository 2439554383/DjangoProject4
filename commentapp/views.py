import json
import subprocess
import tempfile
import uuid
from io import BytesIO
import pymupdf
from docx import Document
from django.http import StreamingHttpResponse
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.db import connection
from PIL import Image
import os
import re
from openai import OpenAI
import requests
import base64
from lxml import etree
import sys
import time


# Face swap utilities
from commentapp.magic.face import load_models, file_to_cv2_image, swap_face_from_cv2


# =========================
# Helper functions (renamed & simplified)
# =========================

def clone_voice_sync(audio, text):
    try:
        url = 'https://aivoiceclonefree.com/api/instant/clone-sync'
        api_key = 'sk-DIFcSSOXhFYwweBgPr3VEIF9CXhIslW1'
        req_data = {'text': text, 'api_key': api_key, 'type': 2}
        files = {'audio': audio}
        resp = requests.post(url, data=req_data, files=files)
        data = json.loads(resp.text)
        return data.get('audio_url')
    except Exception as exc:
        print(f"clone_voice_sync error: {exc}")
        return None


def strip_markdown(text):
    text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[\-\*]\s*", "", text, flags=re.MULTILINE)
    return text.strip()


def fetch_image_b64(url):
    # 添加URL验证
    if not url or url == 'None' or not isinstance(url, str):
        print(f"fetch_image_b64: 无效的URL: {url}")
        return None
    try:
        resp = requests.get(url, timeout=10)  # 添加10秒超时
        if resp.status_code == 200:
            return base64.b64encode(resp.content).decode('utf-8')
        print(f"fetch_image_b64: error status {resp.status_code}")
        return None
    except Exception as exc:
        print(f"fetch_image_b64: 请求失败: {exc}")
        return None


def encode_image_b64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')


def compress_image_to_b64(file, max_size_kb=1024):
    img = Image.open(file)
    buffer = BytesIO()
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    quality = 85
    while True:
        buffer.seek(0)
        buffer.truncate()
        img.save(buffer, format="JPEG", quality=quality)
        size_kb = buffer.tell() / 1024
        if size_kb <= max_size_kb or quality <= 20:
            break
        quality -= 5
    return base64.b64encode(buffer.getvalue()).decode('utf-8')


def gen_image_text_stream(base64_image, text):
    try:
        print(f"[AI] 开始调用VLM模型，图片大小: {len(base64_image) if base64_image else 0} bytes")
        client = OpenAI(api_key="sk-hnrdyinxtiweniixaanaydjbofjwxacqbdmybgcpuqzuzznn",
                        base_url="https://api.siliconflow.cn/v1")
        response = client.chat.completions.create(
            model="Qwen/Qwen3-Omni-30B-A3B-Instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        },
                        {"type": "text", "text": f"{text}"}
                    ]
                }
            ],
            temperature=0.3,
            stream=True,
            max_tokens=1000,  # 限制生成长度避免卡顿
            timeout=180  # 添加超时
        )
        print("[AI] VLM模型调用成功")
        return response
    except Exception as exc:
        print(f"[AI错误] VLM模型调用失败: {exc}")
        import traceback
        traceback.print_exc()
        return None


def sanitize_filename(name):
    return "".join(re.findall(r'[\u4e00-\u9fff\w -]', name)).strip()


def stream_chunks(response):
    full_reply = ""
    first = True
    chunk_count = 0
    try:
        print("[流式输出] 开始接收数据")
        for chunk in response:
            chunk_count += 1
            if chunk_count == 1:
                print("[流式输出] 第一个数据块到达")
            if hasattr(chunk, 'choices') and len(chunk.choices) > 0 and hasattr(chunk.choices[0], 'delta'):
                delta = chunk.choices[0].delta
                if hasattr(delta, 'content') and delta.content:
                    part = delta.content
                    if first:
                        part = "（Ai生成）" + part
                        first = False
                    print(f"[流式输出]: {part}...")
                    full_reply += part
                    yield part
        print(f"[流式输出] 完成，共 {chunk_count} 个数据块，总长度: {len(full_reply)} 字符")
    except Exception as exc:
        print(f"[流式输出错误] {exc}")
        import traceback
        traceback.print_exc()
        yield f"\n\n[错误: {str(exc)}]"


def read_docx_text(file_path):
    try:
        doc = Document(file_path)
        contents = "\n".join([para.text for para in doc.paragraphs])
        return contents
    except Exception as exc:
        print(exc)
        return f"读取 Word 出错：{str(exc)}"


def read_pdf_text(file_path):
    try:
        doc = pymupdf.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as exc:
        print(exc)
        return f"读取 Pdf 出错：{str(exc)}"


def pdf_to_markdown_with_images(file_path):
    doc = pymupdf.open(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "pdf_images")
    os.makedirs(image_dir, exist_ok=True)
    host = "http://139.196.235.10:8005"
    md_text = ""
    image_count = 0
    for page_num in range(len(doc)):
        page = doc[page_num]
        text_blocks = sorted(page.get_text("blocks"), key=lambda b: (b[1], b[0]))
        for block in text_blocks:
            block_text = block[4].strip()
            if not block_text:
                continue
            if len(block_text) < 50 and block_text.endswith(":"):
                md_text += f"\n## {block_text}\n\n"
            else:
                md_text += f"  {block_text}\n\n"
        images = page.get_images(full=True)
        for _, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_count += 1
            image_name = f"pdf_image_{page_num+1}_{image_count}.{image_ext}"
            image_path = os.path.join(image_dir, image_name)
            with open(image_path, "wb") as f:
                f.write(image_bytes)
            md_text += f"![image]({host}/media/pdf_images/{image_name})\n\n"
    doc.close()
    return md_text


def docx_to_markdown_with_images_in_order(file_path):
    doc = Document(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "doc_images")
    os.makedirs(image_dir, exist_ok=True)
    host = "http://139.196.235.10:8005"
    md_text = ""
    image_count = 0
    body = doc._element.body
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    }
    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag == "p":
            blips = child.findall(".//a:blip", namespaces=nsmap)
            if blips:
                for blip in blips:
                    rId = blip.get(f"{{{nsmap['r']}}}embed")
                    image_part = doc.part.related_parts[rId]
                    image_data = image_part.blob
                    image_name = f"image_{image_count + 1}.png"
                    image_path = os.path.join(image_dir, image_name)
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    md_text += f"![image]({host}/media/doc_images/{image_name})\n\n"
                    image_count += 1
                continue
            paragraph = None
            for p in doc.paragraphs:
                if p._p == child:
                    paragraph = p
                    break
            if paragraph:
                style = paragraph.style.name.lower()
                if "heading" in style:
                    level = style.replace("heading ", "")
                    md_text += f"\n{'#' * int(level)} {paragraph.text.strip()}\n\n"
                elif paragraph.text.strip().startswith(("-", "*", "•")):
                    md_text += f"- {paragraph.text.strip().lstrip('-•*')}\n\n"
                else:
                    line = "  "
                    for run in paragraph.runs:
                        t = run.text.replace('\n', ' ').strip()
                        if not t:
                            continue
                        if run.bold and run.italic:
                            line += f"***{t}***"
                        elif run.bold:
                            line += f"**{t}**"
                        elif run.italic:
                            line += f"*{t}*"
                        else:
                            line += t
                    md_text += line + "\n\n"
    return md_text


def transcribe_audio(tmp_audio):
    url = "https://api.siliconflow.cn/v1/audio/transcriptions"
    token = "sk-hnrdyinxtiweniixaanaydjbofjwxacqbdmybgcpuqzuzznn"
    headers = {"Authorization": f"Bearer {token}"}
    model = {"model": "FunAudioLLM/SenseVoiceSmall"}
    
    # 判断tmp_audio是文件对象还是路径字符串
    if isinstance(tmp_audio, str):
        audio_path = tmp_audio
    else:
        audio_path = tmp_audio.name
    
    print(f"[转录] 音频文件路径: {audio_path}")
    print(f"[转录] 文件大小: {os.path.getsize(audio_path) if os.path.exists(audio_path) else '文件不存在'} bytes")
    
    try:
        with open(audio_path, "rb") as f:
            files = {"file": (os.path.basename(audio_path), f, "audio/wav")}
            print(f"[转录] 开始上传音频到API...")
            resp = requests.post(url, headers=headers, data=model, files=files, timeout=120)
            
            if resp.status_code != 200:
                print(f"[转录错误] API返回状态码: {resp.status_code}")
                return ""
            
            result = json.loads(resp.text)
            text = result.get('text', '')
            print(f"[转录完成] 提取的文本: {text[:100] if text else '空'}...")
            return text
    except Exception as exc:
        print(f"[转录错误] 发生异常: {exc}")
        import traceback
        traceback.print_exc()
        return ""


def fetch_video_parse(video_url):
    if "douyin" in video_url:
        platform = "抖音"
        url = "http://api.moreapi.cn/api/douyin/aweme_detail"
    elif "kuaishou" in video_url:
        platform = "快手"
        url = "http://api.moreapi.cn/api/ks/aweme_detail"
    elif "xhslink" in video_url:
        platform = "小红书"
        url = "http://api.moreapi.cn/api/xhs/note_detail"
    elif "toutiao" in video_url:
        platform = "头条"
        url = "http://api.moreapi.cn/api/toutiao/aweme_detail_v2"
    else:
        platform = ""
        url = ""
    payload = json.dumps({"aweme_id": "", "share_text": video_url, "proxy": ""})
    headers = {
        "Authorization": "Bearer O1Y4f9r8sbNdbSqzmpb5MUk3jMS98Hs6exTLosz8bYK0SQyyiQS6nlV2kDDVMghX",
        'Content-Type': 'application/json'
    }
    try:
        resp = requests.request("POST", url, headers=headers, data=payload)
        data = json.loads(resp.text)
        return data, platform
    except Exception as exc:
        print(exc)
        return None, platform


def gen_text_stream(text):
    try:
        print(f"[AI] 开始调用纯文本模型: deepseek-ai/DeepSeek-V2.5")
        print(f"[AI] 文本长度: {len(text)} 字符")
        client = OpenAI(api_key="sk-hnrdyinxtiweniixaanaydjbofjwxacqbdmybgcpuqzuzznn",
                        base_url="https://api.siliconflow.cn/v1")
        response = client.chat.completions.create(
            model='deepseek-ai/DeepSeek-V2.5',
            messages=[{'role': 'user', 'content': f"{text}"}],
            stream=True,
            temperature=0.3,
            max_tokens=1000,  # 限制生成长度
            timeout=180  # 添加超时
        )
        print("[AI] 纯文本模型调用成功")
        return response
    except Exception as exc:
        print(f"[AI错误] 纯文本模型调用失败: {exc}")
        import traceback
        traceback.print_exc()
        return None


# =========================
# Small utilities for complex views
# =========================

def extract_url(text):
    pattern = r'https?://[^\s，。、！？“”‘’<>（）【】()"\']+'
    match = re.search(pattern, text)
    return match.group() if match else None


def select_media_from_moreapi(data, platform):  # 从 moreapi 结构中提取媒体信息
    downloadable_url = ""  # 可下载直链
    cover_image_url = ""  # 封面图直链
    video_title = ""  # 视频标题
    if platform == '抖音':  # 抖音解析
        url_list = data.get('data', {}).get('aweme_detail', {}).get('video', {}).get('play_addr', {}).get('url_list', [])  # 视频直链列表
        cover_list = data.get('data', {}).get('aweme_detail', {}).get('video', {}).get('cover', {}).get('url_list', [])  # 封面直链列表
        downloadable_url = url_list[2] if len(url_list) > 2 else ""  # 选用较高清直链
        cover_image_url = cover_list[1] if len(cover_list) > 1 else ""  # 选用较高清封面
        title_name = data.get('data', {}).get('aweme_detail', {}).get('desc', '')  # 原始标题
        video_title = sanitize_filename(title_name)  # 标题清洗
    elif platform == '小红书':  # 小红书解析
        note_card = data.get('data', {}).get('response_body', {}).get('data', {}).get('items', [{}])[0].get('note_card', {})  # 笔记卡片
        downloadable_url = note_card.get('video', {}).get('media', {}).get('stream', {}).get('h264', [{}])[0].get('master_url', '')  # 视频直链
        image_list = note_card.get('image_list', [{}])[0].get('info_list', [{}])  # 图片列表
        cover_image_url = image_list[0].get('url', '') if image_list else ''  # 封面直链
        video_title = sanitize_filename(note_card.get('title', ''))  # 标题清洗
    elif platform == '快手':  # 快手解析
        ks_data = data.get('data', [{}])[0]  # 结构根
        downloadable_url = ks_data.get('manifest', {}).get('adaptationSet', [{}])[0].get('representation', [{}])[0].get('url', '')  # 视频直链
        cover_image_url = ks_data.get('coverUrls', [{}])[0].get('url', '')  # 封面直链
        video_title = sanitize_filename(ks_data.get('caption', ''))  # 标题清洗
    elif platform == '头条':  # 头条解析
        tt_video = data.get('data', {}).get('data', {}).get('video', {})  # 视频节点
        url_list = tt_video.get('play_addr', {}).get('url_list', [])  # 视频直链列表
        cover_list = tt_video.get('origin_cover', {}).get('url_list', [])  # 封面直链列表
        downloadable_url = url_list[0] if url_list else ""  # 选第一个直链
        cover_image_url = cover_list[0] if cover_list else ""  # 选第一个封面
        video_title = sanitize_filename(data.get('data', {}).get('data', {}).get('title', ''))  # 标题清洗
    return downloadable_url, cover_image_url, video_title  # 返回三要素


def fallback_parse_api(video_url):  # 第二方案兜底解析
    try:
        resp = requests.get(f"https://api.yyy001.com/api/videoparse?url={video_url}", timeout=7)  # 调用兜底解析API
        if resp.status_code == 200:  # 成功
            data = resp.json()  # JSON解析
            downloadable_url = data.get('data', {}).get('url', None)  # 可下载直链
            cover_image_url = data.get('data', {}).get('cover')  # 封面直链
            video_title = sanitize_filename(data.get('data', {}).get('title', ''))  # 标题清洗
            return downloadable_url, cover_image_url, video_title  # 返回三要素
        return None, None, None  # 状态码非200
    except Exception as exc:
        print("fallback_parse_api error:", exc)  # 打印错误
        return None, None, None  # 异常兜底


def download_audio_30s_ffmpeg(video_url):
    with tempfile.NamedTemporaryFile(suffix=".wav") as tmp_audio:
        cmd = [
            "ffmpeg", "-y",
            "-t", "30",
            "-i", video_url,
            "-vn", "-ar", "16000", "-ac", "1", "-f", "wav",
            tmp_audio.name
        ]
        subprocess.run(cmd, check=True)
        return tmp_audio


COMMENT_PROMPT_TEMPLATES = {
    "幽默": "以轻松诙谐的口吻描写，制造包袱或反转但保持友善，收尾要有记忆点。把幽默点落在视频核心内容上，控制在约{number}字。",
    "干货": "给出1-2个具体观点或方法，突出实用价值，语言精练有逻辑，不空泛，保持约{number}字。",
    "热词玩梗": "自然穿插当下热门梗或网络热词，但不要堆砌，保持与视频内容紧密相关，约{number}字。",
    "散文诗歌": "用散文化诗意语言描绘观感，句式富有节奏感与画面感，情绪细腻饱满，控制在约{number}字左右。",
    "咨询": "以真诚求教的语气提出1-2个精准问题，表现出对细节的关注和行动意愿，字数约{number}字。",
    "感同身受": "表达强烈共鸣，引用音频或画面中的细节作为共情支撑，让读者感到真挚，约{number}字。",
    "李白风格的唐诗": "创作一首仿李白风骨的七言古风诗，意象豪迈洒脱，可借酒、月、江山等意象，灵活分为约{lines}句，总字数接近{number}字。",
    "点赞": "直接表达喜爱与赞赏，点明喜欢的原因或细节，语气热情真诚，约{number}字。",
    "暖心鼓励": "用温暖有力量的话鼓励作者或主角，体现理解与支持，给出积极期待，字数约{number}字。",
    "董宇辉式小作文": "采用真诚温暖、有画面感的文风，引用生活经验或文化典故，引导读者共鸣，层次分明，约{number}字。",
    "宋词": "创作一首符合宋词婉约或豪放格调的小词，结构完整，词意连贯柔婉或豪迈，总字数贴近{number}字。",
    "高情商": "保持尊重与体贴，以委婉方式表达观点或建议，兼顾对方面子和感受，约{number}字。",
    "七言绝句": "写成约{lines}句的七言体诗，保持起承转合与平仄韵律，句句紧扣视频主题，总字数约{number}字。",
    "神评论": "用一句或短段极具洞察力的话，巧妙点题或反转，引人点赞转发，约{number}字。",
    "抒情七言绝句": "以细腻情感创作约{lines}句七言体诗，情绪层层递进，总字数约{number}字。",
    "咏物诗七言绝句": "围绕视频中的关键事物写约{lines}句七言体诗，借物抒怀，总字数控制在{number}字左右。",
    "叙事七言绝句": "用七言句式写成约{lines}句的小叙事诗，讲清事件脉络，总字数贴近{number}字。",
    "讨论七言绝句": "在约{lines}句七言体诗中融入观点或疑问，引导互动，总字数约{number}字。",
    "山水田园七言绝句": "描摹山水田园意境，写成约{lines}句七言体诗，借景抒情，总字数约{number}字。",
    "边塞七言绝句": "以雄浑苍凉的语调创作约{lines}句七言体诗，展现边塞豪情，总字数约{number}字。",
    "婉约七言绝句": "创作约{lines}句七言体诗，语调婉约柔美，情感内敛含蓄，总字数约{number}字。",
    "豪放七言绝句": "用七言句式创作约{lines}句豪放之诗，节奏明快有力量，总字数约{number}字。",
    "加油": "以热血积极的语气为作者或主角打气，可给出具体期待或目标，约{number}字。",
    "支持": "明确表态支持，说明理由或未来行动，语气坚定友善，控制在{number}字左右。",
    "同意": "先简洁复述对方观点亮点，再补充自己的认同理由或延伸思考，约{number}字。",
    "羡慕": "表达由衷羡慕，点出最触动你的细节，并带出自己的想法或愿望，约{number}字。",
    "向往": "描绘你对这种生活或体验的向往，结合标题与音频细节，展现期待与计划，约{number}字。",
    "咨询类默认": "保持礼貌谦逊，提出一到两个核心问题，明确你想进一步了解的关键点，约{number}字。",
    "默认": "以真实用户口吻结合视频细节和音频信息输出评论，语言自然顺畅，观点清晰，控制在{number}字左右。"
}

POETRY_STYLES_NEED_LINES = {
    "李白风格的唐诗",
    "七言绝句",
    "抒情七言绝句",
    "咏物诗七言绝句",
    "叙事七言绝句",
    "讨论七言绝句",
    "山水田园七言绝句",
    "边塞七言绝句",
    "婉约七言绝句",
    "豪放七言绝句",
}


def normalize_comment_type(raw_value):
    if isinstance(raw_value, list):
        return raw_value[0].strip() if raw_value else ""
    if raw_value is None:
        return ""
    return str(raw_value).strip()


def build_comment_prompt(audio_text, title, number, comment_type):
    number = number or 50
    comment_type = comment_type or ""
    audio_text = audio_text.strip() if audio_text else ""
    title = title.strip() if title else ""
    safe_title = title or "未提供标题"
    safe_audio = audio_text or "（音频内容为空或未识别）"
    base_context = (
        f"视频标题：《{safe_title}》。\n"
        f"音频文本摘录：{safe_audio}。\n"
        f"请基于以上标题与音频内容，创作约{number}字的短视频评论（允许上下浮动2-3个字）"
        f"输出内容直接输出评论内容，而不是出现标题：或者评论：这种格式，直接给我评论内容"
    )
    template = COMMENT_PROMPT_TEMPLATES.get(comment_type)
    lines = 0
    if comment_type in POETRY_STYLES_NEED_LINES:
        try:
            lines = max(4, round(int(number) / 7)) if number else 4
        except (TypeError, ValueError):
            lines = 4
    if not template:
        if comment_type:
            if "咨询" in comment_type:
                template = COMMENT_PROMPT_TEMPLATES.get("咨询类默认")
                detail_instruction = template.format(
                    number=number,
                    title=safe_title,
                    audio_text=safe_audio,
                    lines=lines
                )
            else:
                detail_instruction = (
                    f"请以“{comment_type}”风格创作评论，语言自然流畅，突出该类型的典型特征，控制在约{number}字。"
                )
        else:
            template = COMMENT_PROMPT_TEMPLATES.get("默认")
            detail_instruction = template.format(
                number=number,
                title=safe_title,
                audio_text=safe_audio,
                lines=lines
            )
    else:
        detail_instruction = template.format(
            number=number,
            title=safe_title,
            audio_text=safe_audio,
            lines=lines
        )
    if comment_type and (
        comment_type in POETRY_STYLES_NEED_LINES
        or any(keyword in comment_type for keyword in ["宋词", "词", "绝句", "唐诗"])
    ):
        detail_instruction += (
            "\n请仅输出符合要求的诗词正文，不要添加标题、备注、解释或其他白话内容，更不要出现提示语。"
        )
    prompt = f"{base_context}\n{detail_instruction}"
    print(f"生成评论提示词类型: {comment_type}，提示内容: {prompt}...")
    return prompt



def image_to_base64(path):
    with open(path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')


# =========================
# Views (external names kept; internals refactored)
# =========================

@csrf_exempt
def get_comment(request):  # 生成短视频评论
    if request.method != 'POST':  # 仅支持POST
        return JsonResponse({"result": False}, status=500)  # 方法不符
    data = json.loads(request.body)  # 读取JSON体
    number = data.get('count')  # 评论字数目标
    shared_text = data.get('url', '')  # 分享文本（可能包含链接）
    comment_type = normalize_comment_type(
        data.get("selecttext_list")
    ) or normalize_comment_type(
        data.get("comment_type")
    ) or normalize_comment_type(
        data.get("prompt_type")
    ) or normalize_comment_type(
        data.get("selecttext")
    )
    video_url = extract_url(shared_text)  # 提取视频URL
    if not video_url:  # 无URL
        return JsonResponse({"result": False}, status=500)  # 返回失败
    try:
        parse_data, platform = fetch_video_parse(video_url)  # 第一方案解析
        if parse_data:  # 解析成功
            downloadable_url, cover_image_url, video_title = select_media_from_moreapi(parse_data, platform)  # 提取三要素
        else:  # 解析失败
            downloadable_url, cover_image_url, video_title = None, None, None  # 置空
    except Exception as exc:  # 异常进入兜底
        print("fetch_video_parse error, fallback:", exc)  # 打印错误
        downloadable_url, cover_image_url, video_title = fallback_parse_api(video_url)  # 兜底解析
    if downloadable_url:  # 有可下载直链
        tmp_audio = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        audio_path = None
        try:
            print("=" * 60)
            print("步骤1: 提取音频")
            print("=" * 60)
            cmd = [  # ffmpeg命令
                "ffmpeg", "-y", "-t", "30", "-i", downloadable_url,
                "-vn", "-ar", "16000", "-ac", "1", "-f", "wav", tmp_audio.name
            ]
            subprocess.run(cmd, check=True)  # 执行抽音频
            tmp_audio.flush()
            audio_path = tmp_audio.name
            tmp_audio.close()
            
            print("=" * 60)
            print("步骤2: 音频转文本")
            print("=" * 60)
            audio_text = transcribe_audio(audio_path)  # 语音转文本
            print(f"音频文本: {audio_text[:100] if audio_text else '空'}...")
            
            print("=" * 60)
            print("步骤3: 构造提示词并获取封面图")
            print("=" * 60)
            prompt = build_comment_prompt(audio_text, video_title, number, comment_type)  # 构造提示词
            
            base64_image = None
            if cover_image_url and cover_image_url != "None":
                base64_image = fetch_image_b64(cover_image_url)  # 获取封面Base64
                if base64_image:
                    print(f"封面图获取成功，大小: {len(base64_image)} bytes")
                else:
                    print("封面图获取失败")
            
            print("=" * 60)
            print("步骤4: 调用AI生成评论")
            print("=" * 60)
            
            # 根据是否有封面图选择模型
            if base64_image:
                # print("使用VLM图文模型")
                # response_data = gen_image_text_stream(base64_image, prompt)
                print("有封面图但是直接使用文本模式")
                # 修改提示词去掉图片相关内容
                prompt_text = prompt
                response_data = gen_text_stream(prompt_text)
            else:
                print("无封面图，使用纯文本模型")
                # 修改提示词去掉图片相关内容
                prompt_text = prompt
                response_data = gen_text_stream(prompt_text)
            if not response_data:
                return JsonResponse({"result": False, "msg": "AI生成失败"}, status=500)
            
            print("=" * 60)
            print("步骤5: 开始流式返回")
            print("=" * 60)
            return StreamingHttpResponse(stream_chunks(response_data), content_type='text/plain')
        except Exception as exc:
            print(f"处理失败: {exc}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"result": False, "msg": str(exc)}, status=500)
        finally:
            if audio_path and os.path.exists(audio_path):
                try:
                    os.unlink(audio_path)
                except:
                    pass
    else:  # 没有视频直链，仅用封面与标题生成
        print("没有视频直链，仅用封面与标题生成评论")
        try:
            # 构造提示词
            prompt = build_comment_prompt("", video_title, number, comment_type)
            
            # 获取封面图
            base64_image = None
            if cover_image_url and cover_image_url != "None":
                base64_image = fetch_image_b64(cover_image_url)
                if base64_image:
                    print(f"封面图获取成功，大小: {len(base64_image)} bytes")
            print(f"最终提示词{prompt}")
            # 根据是否有封面图选择模型
            if base64_image:
                print("有封面图，使用纯文本模型")
                response_data = gen_text_stream(prompt)
                # response_data = gen_image_text_stream(base64_image, prompt)
            else:
                print("无封面图，使用纯文本模型")
                response_data = gen_text_stream(prompt)
            
            if not response_data:
                return JsonResponse({"result": False, "msg": "AI生成失败"}, status=500)
            
            return StreamingHttpResponse(stream_chunks(response_data), content_type='text/plain')
        except Exception as exc:
            print(f"处理失败: {exc}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"result": False, "msg": str(exc)}, status=500)


@csrf_exempt
def get_listmodel(request):
    if request.method != 'POST':
        return JsonResponse({"status": False}, status=500)
    data = json.loads(request.body)
    code = data.get("code")
    with connection.cursor() as cursor:
        cursor.execute('select * from comment_type where user_code = %s', (code,))
        type_list = [
            (row[0], row[1], bool(row[2]), bool(row[3]))
            for row in cursor.fetchall()
        ]
    with connection.cursor() as cursor:
        cursor.execute('select * from comment_type where isset=1 and user_code = %s', (code,))
        type_overlay_list = [
            (row[0], row[1], bool(row[2]), bool(row[3]))
            for row in cursor.fetchall()
        ]
    return JsonResponse({"status": True, "type_list": type_list, "type_overlay_list": type_overlay_list})


@csrf_exempt
def switch_isset(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    code = data.get("code")
    with connection.cursor() as cursor:
        for i in range(len(data.get('list_type', []))):
            cursor.execute(
                'update comment_type set isset=%s where user_code = %s and type_name = %s',
                (data['list_type'][i][2], code, data['list_type'][i][1])
            )
    return JsonResponse({"data": True})


@csrf_exempt
def get_hospitallist(request):
    if request.method != 'GET':
        return JsonResponse({"data": False}, status=500)
    file_list = []
    path_file = '/www/wwwroot/DjangoProject4/data/hospital/'
    for file in os.listdir(path_file):
        if file.endswith(".xlsx"):
            file_list.append({"name": file, "path": f"http://139.196.235.10:8005/media/hospital/{file}"})
    return JsonResponse({"data": file_list})


def get_filelist(request):
    if request.method != 'GET':
        return JsonResponse({"data": False}, status=500)
    folder_dict = {}
    file_path = '/www/wwwroot/DjangoProject4/data/file/'
    for folder in os.listdir(file_path):
        file_list = []
        folder_path = os.path.join(file_path, folder)
        if os.path.isdir(folder_path) and len(os.listdir(folder_path)):
            for file in os.listdir(folder_path):
                if file.endswith(".pdf"):
                    path = f"{folder_path}/{file}"
                    name = os.path.splitext(file)[0]
                    content = pdf_to_markdown_with_images(path)
                    file_list.append({"name": name, "content": content})
                elif file.endswith(".docx"):
                    path = f"{folder_path}/{file}"
                    name = os.path.splitext(file)[0]
                    content = docx_to_markdown_with_images_in_order(path)
                    file_list.append({"name": name, "content": content})
        folder_dict[folder] = file_list
    return JsonResponse({"data": folder_dict})


@csrf_exempt
def get_image(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    file = request.FILES['file']
    text = request.POST['text']
    typ = request.POST['type']
    image_base64 = compress_image_to_b64(file, max_size_kb=1024)
    if typ == '生成菜单':
        prompt = f'根据 {text} 这几个要求和图中的食材，智能生成菜谱，输出结果只需要菜谱名和菜谱流程'
    elif typ == '查热量':
        prompt = f'根据 {text} 这几个要求和图中的食物智能识别种类，精准估算热量与营养成分，输出结果只需要输出食物种类和名称，热量还有营养成分'
    else:
        prompt = text
    response = gen_image_text_stream(image_base64, prompt)
    if not response:
        return JsonResponse({"data": False}, status=500)
    return StreamingHttpResponse(stream_chunks(response), content_type='text/plain')


@csrf_exempt
def get_text(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    text = data.get('text')
    typ = data.get('type')
    if typ == '姓名打分':
        prompt = f'根据 {text} 这个姓名，分析五行、音律、寓意等多维度，给出权威综合评分'
    elif typ == '起名':
        prompt = f'根据 {text} 这个姓氏进行智能取名，兼顾音义美与吉祥寓意'
    elif typ == '智能助手':
        prompt = f'{text}'
    else:
        prompt = f'{text}'
    response = gen_text_stream(prompt)
    return StreamingHttpResponse(stream_chunks(response), content_type='text/plain')


# @csrf_exempt
# def get_code(request):
#     if request.method != 'POST':
#         return JsonResponse({"data": False})
#     default_types = [
#         '高情商','同意','幽默','支持','提问','感动','暖心','鼓励','加油','反对','质疑','批评',
#         '惊讶','不可思议','羡慕','向往','求解答','召唤','讨论','标记','收藏','干货','有用',
#         '求教程','求链接','分享经验','补充信息','热词玩梗','简短有力','神评论','表达喜爱','催更',
#         '夸赞博主','价格咨询','产品细节追问','真人测评诉求','竞品对比','售后担忧','场景化需求',
#         '追问原理求资料','周星驰式','梁朝伟式','预言','赞美','董宇辉式小作文','七言绝句',
#         '散文诗歌','唐诗','宋词','歌词'
#     ]
#     data = json.loads(request.body)
#     code = data.get('code')
#     with connection.cursor() as cursor:
#         cursor.execute('select * from user where code = %s', (code,))
#         isexist = cursor.fetchall()
#         if isexist:
#             return JsonResponse({"data": False})
#         cursor.execute('insert into user(code) values (%s)', (code,))
#         for item in default_types:
#             cursor.execute(
#                 'insert into comment_type(type_name,isset,ischeck,user_code) values (%s,1,0,%s)',
#                 (item, code)
#             )
#     return JsonResponse({"data": True})


@csrf_exempt
def post_content(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    code = data.get('code')
    content = data.get('content')
    if content:
        with connection.cursor() as cursor:
            cursor.execute('insert into content(content,user_code) values (%s,%s)', (content, code))
        with connection.cursor() as cursor:
            cursor.execute('select content from content where user_code=%s', (code,))
            result = [row[0] for row in cursor.fetchall()]
        return JsonResponse({"data": result})
    else:
        with connection.cursor() as cursor:
            cursor.execute('select content from content where user_code=%s order by id desc', (code,))
            result = [row[0] for row in cursor.fetchall()]
        return JsonResponse({"data": result})


@csrf_exempt
def post_aiface(request):  # 建议项目启动时加载一次
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    old_image = request.FILES['old_image']
    face_image = request.FILES['face_image']
    load_models()
    input_img = file_to_cv2_image(old_image)
    face_img = file_to_cv2_image(face_image)
    result_base64 = swap_face_from_cv2(input_img, face_img)
    if result_base64:
        return JsonResponse({'data': {"code": 200, "data": result_base64}})
    return JsonResponse({'error': '换脸失败'}, status=500)


@csrf_exempt
def get_aiimage(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    try:
        data = json.loads(request.body)
        text = data.get('text')
        url = "https://api.siliconflow.cn/v1/images/generations"
        payload = {
            "model": "Kwai-Kolors/Kolors",
            "prompt": text,
            "image_size": "1024x1024",
            "batch_size": 1,
            "seed": 4999999999,
            "num_inference_steps": 20,
            "guidance_scale": 7.5,
        }
        headers = {
            "Authorization": "Bearer sk-hnrdyinxtiweniixaanaydjbofjwxacqbdmybgcpuqzuzznn",
            "Content-Type": "application/json"
        }
        resp = requests.request("POST", url, json=payload, headers=headers)
        result = json.loads(resp.text)
        image_url = result["data"][0]["url"]
        return JsonResponse({"data": image_url})
    except Exception as exc:
        print(exc)
        return JsonResponse({"data": False}, status=500)


@csrf_exempt
def get_unmarkvideo(request):  # 获取无水印视频直链
    if request.method != 'POST':  # 仅支持POST
        return JsonResponse({"result": False}, status=500)  # 方法不符
    try:
        data = json.loads(request.body)  # 读取JSON体
        text = data.get('text', '')  # 分享文本
        video_url = extract_url(text)  # 提取链接
        if not video_url:  # 无URL
            return JsonResponse({"result": False}, status=500)  # 返回失败
        try:  # 第一方案
            parse_data, platform = fetch_video_parse(video_url)  # 解析平台
            if not parse_data:  # 解析失败
                raise RuntimeError("primary parse failed")  # 抛出异常走兜底
            downloadable_url, cover_image_url, video_title = select_media_from_moreapi(parse_data, platform)  # 提取直链
        except Exception as exc:  # 兜底方案
            print("primary parse failed:", exc)  # 打印错误
            downloadable_url, cover_image_url, video_title = fallback_parse_api(video_url)  # 兜底解析
            if not downloadable_url:  # 仍失败
                return JsonResponse({"result": False}, status=500)  # 返回失败
        return JsonResponse({"data": downloadable_url})  # 返回直链
    except Exception as exc:  # 总异常兜底
        print(exc)  # 打印错误
        return JsonResponse({"result": False}, status=500)  # 返回失败


@csrf_exempt
def post_audio(request):
    if request.method != 'POST':
        return JsonResponse({'error': '获取失败'}, status=500)
    try:
        text = request.POST.get('text')
        file = request.FILES.get('file')
        result = clone_voice_sync(file, text)
        if not result:
            return JsonResponse({'error': '获取失败'}, status=500)
        return JsonResponse({'data': result})
    except Exception as exc:
        print(exc)
        return JsonResponse({'error': '获取失败'}, status=500)


@csrf_exempt
def voice_list(request):
    if request.method != 'GET':
        return JsonResponse({"status": 0, "msg": "仅支持GET请求", "data": []})
    try:
        base_path = "/www/wwwroot/DjangoProject4/data/voice"
        domain = request.build_absolute_uri('/')[:-1]
        media_url_prefix = "/media/voice/"
        if not os.path.exists(base_path):
            return JsonResponse({"status": 0, "msg": "目录不存在", "data": []})
        audio_exts = (".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a")
        files = [f for f in os.listdir(base_path) if f.lower().endswith(audio_exts)]
        urls = [domain + media_url_prefix + f for f in files]
        return JsonResponse({"status": 1, "msg": "success", "data": urls})
    except Exception as exc:
        return JsonResponse({"status": 0, "msg": str(exc), "data": []})


def call_voice_separation_api(audio_url, language="zh"):  # 调用302.ai人声分离API
    """
    调用302.ai的人声分离API
    :param audio_url: 音频链接
    :param language: 音频语言，默认中文
    :return: 返回task_id或None
    """
    try:
        api_url = "https://api.302.ai/302/vt/subtitle/extract"  # 302.ai API端点
        headers = {
            "Authorization": "Bearer sk-ilbIeyYIgSlY63gemjvJvKwsFwyYKyE7MdGa6b4SHdvL0aZl",  # 需要替换为实际的API密钥
            "Content-Type": "application/json"
        }
        payload = {
            "audio_url": audio_url,  # 音频链接
            "language": language,  # 语言
            "demucs": True,  # 开启人声分离
            "is_only_demucs": True  # 只分离人声，不转录
        }
        
        print(f"正在调用人声分离API，音频URL: {audio_url}")  # 调试日志
        response = requests.post(api_url, json=payload, headers=headers, timeout=30)  # 发送请求
        
        if response.status_code == 200:  # 请求成功
            result = response.json()  # 解析JSON响应
            print(f"人声分离API调用成功，结果: {result}")  # 调试日志
            task_id = result.get('task_id')  # 获取任务ID
            print(f"人声分离API调用成功，任务ID: {task_id}")  # 调试日志
            return task_id  # 返回任务ID
        else:
            print(f"人声分离API调用失败，状态码: {response.status_code}, 响应: {response.text}")  # 错误日志
            return None  # 返回失败
            
    except Exception as exc:
        print(f"人声分离API调用异常: {exc}")  # 异常日志
        return None  # 返回失败


def get_separated_voice_url(task_id):  # 根据任务ID获取分离后的人声URL
    """
    根据任务ID查询分离后的人声URL
    调用302.ai的任务状态查询API
    :param task_id: 任务ID
    :return: 返回字典 {"status": "success/processing/queue/fail", "voice_url": "url或None"}
    """
    try:
        # 根据302.ai官方文档实现查询接口
        query_url = f"https://api.302.ai/302/vt/tasks/subtitle/{task_id}"  # API端点
        headers = {
            "Authorization": "Bearer sk-ilbIeyYIgSlY63gemjvJvKwsFwyYKyE7MdGa6b4SHdvL0aZl",  # API密钥
        }
        
        print(f"正在查询任务状态，任务ID: {task_id}")  # 调试日志
        
        # 发送GET请求查询任务状态
        response = requests.get(query_url, headers=headers, timeout=30)  # 30秒超时
        
        if response.status_code == 200:  # 请求成功
            result = response.json()  # 解析JSON响应
            status = result.get('status')  # 获取任务状态：queue/processing/success/fail
            
            print(f"任务状态: {status}")  # 调试日志
            print(f"完整响应数据: {json.dumps(result, ensure_ascii=False)}")  # 打印完整响应
            
            if status == 'success':  # 任务成功完成
                # 从result中获取分离后的音频URL
                result_data = result.get('result', {})
                
                print(f"result数据: {json.dumps(result_data, ensure_ascii=False)}")  # 打印result内容
                
                # 根据302.ai实际返回数据，直接获取vocal_audio_url（人声音频）
                voice_url = result_data.get('vocal_audio_url')  # 人声音频URL（主要字段）
                background_url = result_data.get('background_audio_url')  # 背景音频URL
                
                # 如果没有vocal_audio_url，尝试其他可能的字段
                if not voice_url:
                    voice_url = (
                        result_data.get('vocals_url') or  # 人声URL
                        result_data.get('vocal_url') or  # 人声URL
                        result_data.get('voice_url') or  # 音频URL
                        result_data.get('audio_url') or  # 音频URL
                        result_data.get('video_url') or  # 视频URL
                        result_data.get('url')  # 通用URL字段
                    )
                
                print(f"人声分离成功，人声URL: {voice_url}, 背景URL: {background_url}")  # 调试日志
                return {
                    "status": "success", 
                    "voice_url": voice_url,  # 人声音频URL
                    "background_url": background_url,  # 背景音频URL
                    "result": result_data  # 完整结果数据
                }
                
            elif status in ['queue', 'processing']:  # 任务还在处理中
                print(f"任务仍在处理中，状态: {status}")  # 调试日志
                return {"status": status, "voice_url": None}  # 返回处理中状态
                
            elif status == 'fail':  # 任务失败
                print(f"任务失败")  # 调试日志
                return {"status": "fail", "voice_url": None}  # 返回失败状态
            else:
                print(f"未知任务状态: {status}")  # 调试日志
                return {"status": "unknown", "voice_url": None}  # 返回未知状态
        else:
            print(f"查询任务状态失败，状态码: {response.status_code}, 响应: {response.text}")  # 错误日志
            return {"status": "error", "voice_url": None}  # 返回错误状态
        
    except Exception as exc:
        print(f"查询人声分离结果异常: {exc}")  # 异常日志
        return {"status": "error", "voice_url": None}  # 返回失败


@csrf_exempt
def query_voice_separation_task(request):  # 查询人声分离任务状态
    """
    供前端轮询查询人声分离任务状态的接口
    """
    if request.method != 'POST':  # 仅支持POST请求
        return JsonResponse({"status": False, "msg": "仅支持POST请求"}, status=405)  # 方法不允许
    
    try:
        data = json.loads(request.body)  # 解析JSON请求体
        task_id = data.get('task_id', '')  # 获取任务ID
        
        if not task_id:  # 未提供任务ID
            return JsonResponse({"status": False, "msg": "缺少任务ID参数"}, status=400)  # 返回错误
        
        print(f"查询任务状态，任务ID: {task_id}")  # 调试日志
        
        # 调用查询函数
        result = get_separated_voice_url(task_id)  # 查询分离结果
        task_status = result.get('status')  # 获取任务状态
        voice_url = result.get('voice_url')  # 获取人声URL
        background_url = result.get('background_url')  # 获取背景音频URL
        
        if task_status == 'success' and voice_url:  # 成功获取人声URL
            return JsonResponse({  # 返回成功结果
                "status": True,
                "message": "人声分离成功",
                "data": {
                    "task_id": task_id,  # 任务ID
                    "task_status": task_status,  # 任务状态
                    "voice_url": voice_url,  # 人声音频URL（纯人声）
                    "background_url": background_url,  # 背景音频URL（背景音乐）
                    "result_detail": result.get('result', {})  # 详细结果
                }
            })
        elif task_status in ['queue', 'processing']:  # 任务仍在处理中
            return JsonResponse({  # 返回处理中状态
                "status": True,
                "message": f"任务正在处理中，当前状态：{task_status}",
                "data": {
                    "task_id": task_id,  # 任务ID
                    "task_status": task_status,  # 任务状态
                    "voice_url": None  # 人声URL暂未生成
                }
            })
        else:  # 任务失败或错误
            return JsonResponse({  # 返回失败信息
                "status": False,
                "message": f"任务失败，状态：{task_status}",
                "data": {
                    "task_id": task_id,  # 任务ID
                    "task_status": task_status  # 任务状态
                }
            }, status=500)
            
    except Exception as exc:  # 总异常处理
        print(f"查询任务状态异常: {exc}")  # 打印错误
        return JsonResponse({"status": False, "message": f"查询异常: {str(exc)}"}, status=500)  # 返回错误


@csrf_exempt
def extract_voice_from_video(request):  # 从短视频中提取人声
    """
    接收包含短视频链接的文案，解析视频获取音频，然后调用人声分离API
    """
    if request.method != 'POST':  # 仅支持POST请求
        return JsonResponse({"status": False, "msg": "仅支持POST请求"}, status=405)  # 方法不允许
    
    try:
        data = json.loads(request.body)  # 解析JSON请求体
        video_text = data.get('text', '')  # 获取包含视频链接的文案
        language = data.get('language', 'zh')  # 获取语言参数，默认中文
        
        print(f"开始处理视频文案: {video_text}")  # 调试日志
        
        # 第一步：提取视频链接
        video_url = extract_url(video_text)  # 从文案中提取视频URL
        if not video_url:  # 未找到视频链接
            return JsonResponse({"status": False, "msg": "未在文案中找到有效的视频链接"}, status=400)  # 返回错误
        
        print(f"提取到视频链接: {video_url}")  # 调试日志
        
        # 第二步：解析视频信息
        try:  # 尝试第一方案解析
            parse_data, platform = fetch_video_parse(video_url)  # 调用解析API
            if parse_data:  # 解析成功
                downloadable_url, cover_image_url, video_title = select_media_from_moreapi(parse_data, platform)  # 提取媒体信息
            else:  # 解析失败
                downloadable_url, cover_image_url, video_title = None, None, None  # 置空
        except Exception as exc:  # 异常进入兜底
            print(f"第一方案解析失败: {exc}")  # 打印错误
            downloadable_url, cover_image_url, video_title = fallback_parse_api(video_url)  # 调用兜底解析
        
        if not downloadable_url:  # 没有获取到视频链接
            return JsonResponse({"status": False, "msg": "无法解析视频链接，请检查链接是否有效"}, status=400)  # 返回错误
        
        print(f"成功解析视频，直链: {downloadable_url}")  # 调试日志
        
        # 第三步：调用人声分离API
        task_id = call_voice_separation_api(downloadable_url, language)  # 调用人声分离API
        if not task_id:  # API调用失败
            return JsonResponse({"status": False, "msg": "人声分离API调用失败"}, status=500)  # 返回错误
        
        print(f"人声分离任务已提交，任务ID: {task_id}")  # 调试日志
        
        # 第四步：轮询等待任务完成
        max_wait_time = 300  # 最大等待时间：300秒（5分钟）
        poll_interval = 3  # 每3秒查询一次
        elapsed_time = 0  # 已等待时间
        
        while elapsed_time < max_wait_time:  # 在最大等待时间内循环
            result = get_separated_voice_url(task_id)  # 查询分离结果
            task_status = result.get('status')  # 获取任务状态
            voice_url = result.get('voice_url')  # 获取人声URL
            background_url = result.get('background_url')  # 获取背景音频URL
            
            print(f"轮询查询 [{elapsed_time}s]: 任务状态={task_status}, voice_url={'有' if voice_url else '无'}")  # 调试日志
            
            if task_status == 'success':  # 任务成功完成
                print(f"人声分离任务完成，总耗时: {elapsed_time}秒")  # 调试日志
                print(f"人声URL: {voice_url}")  # 调试日志
                print(f"背景URL: {background_url}")  # 调试日志
                
                # 如果没有获取到URL，返回完整的result数据用于调试
                if not voice_url:
                    print(f"警告：任务成功但未获取到音频URL，返回完整结果数据用于调试")  # 警告日志
                    return JsonResponse({  # 返回结果（带调试信息）
                        "status": False,
                        "msg": "任务完成但未找到音频URL，请检查API返回数据",
                        "data": {
                            "task_id": task_id,  # 任务ID
                            "task_status": task_status,  # 任务状态
                            "voice_url": None,  # 人声URL
                            "video_title": video_title,  # 视频标题
                            "platform": platform if 'platform' in locals() else "未知",  # 平台信息
                            "process_time": elapsed_time,  # 处理耗时（秒）
                            "debug_info": result.get('result', {})  # 调试信息：完整的result数据
                        }
                    }, status=500)
                
                # 成功获取到URL
                return JsonResponse({  # 返回成功结果
                    "status": True,
                    "msg": "人声分离成功",
                    "data": {
                        "task_id": task_id,  # 任务ID
                        "task_status": task_status,  # 任务状态
                        "voice_url": voice_url,  # 人声音频URL（纯人声）
                        "background_url": background_url,  # 背景音频URL（背景音乐）
                        "video_title": video_title,  # 视频标题
                        "platform": platform if 'platform' in locals() else "未知",  # 平台信息
                        "process_time": elapsed_time  # 处理耗时（秒）
                    }
                })
            elif task_status == 'fail':  # 任务失败
                print(f"人声分离任务失败")  # 调试日志
                return JsonResponse({  # 返回失败信息
                    "status": False,
                    "msg": "人声分离任务失败",
                    "data": {
                        "task_id": task_id,  # 任务ID
                        "task_status": task_status,  # 任务状态
                        "video_title": video_title,  # 视频标题
                        "platform": platform if 'platform' in locals() else "未知"  # 平台信息
                    }
                }, status=500)
            elif task_status in ['queue', 'processing']:  # 任务仍在处理中
                print(f"任务处理中，等待 {poll_interval} 秒后重试...")  # 调试日志
                time.sleep(poll_interval)  # 等待指定时间
                elapsed_time += poll_interval  # 累加已等待时间
            elif task_status == 'error':  # API返回错误
                print(f"任务状态错误: {task_status}")  # 调试日志
                return JsonResponse({  # 返回错误信息
                    "status": False,
                    "msg": f"查询任务状态失败: {task_status}",
                    "data": {
                        "task_id": task_id,  # 任务ID
                        "task_status": task_status  # 任务状态
                    }
                }, status=500)
            else:  # 其他未知状态，继续等待
                print(f"未知任务状态: {task_status}，继续等待...")  # 调试日志
                time.sleep(poll_interval)  # 等待指定时间
                elapsed_time += poll_interval  # 累加已等待时间
        
        # 超时未完成
        print(f"人声分离任务超时，已等待 {max_wait_time} 秒")  # 调试日志
        return JsonResponse({  # 返回超时信息
            "status": False,
            "msg": f"任务处理超时，请稍后使用任务ID查询结果",
            "data": {
                "task_id": task_id,  # 任务ID
                "task_status": "timeout",  # 超时状态
                "video_title": video_title,  # 视频标题
                "platform": platform if 'platform' in locals() else "未知"  # 平台信息
            }
        }, status=408)  # 408 Request Timeout
            
    except Exception as exc:  # 总异常处理
        print(f"提取人声异常: {exc}")  # 打印错误
        return JsonResponse({"status": False, "msg": f"处理异常: {str(exc)}"}, status=500)  # 返回错误


