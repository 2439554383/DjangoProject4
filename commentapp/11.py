import json
import subprocess
import tempfile
import uuid
from io import BytesIO
import pymupdf
from docx import Document
from django.http import StreamingHttpResponse
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.db import connection
from PIL import Image
import os
import re
from openai import OpenAI
import requests
import base64
from lxml import etree
import sys
import os
import sys
import os


from commentapp.magic.face import load_models, file_to_cv2_image, swap_face_from_cv2

# 添加项目根目录路径
# BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# if BASE_DIR not in sys.path:
#     sys.path.insert(0, BASE_DIR)
#
from magic.face import load_models, swap_face, file_to_cv2_image, swap_face_from_cv2
print("sys.path:", sys.path)
def voice_copy(audio,text):
    try:
        url = 'https://aivoiceclonefree.com/api/instant/clone-sync'
        api_key = 'sk-DIFcSSOXhFYwweBgPr3VEIF9CXhIslW1'
        type = 2
        data = {
            'text': text,
            'api_key': api_key,
            'type': type
        }
        files = {'audio': audio}
        print("开始克隆")
        response = requests.post(url, data=data, files=files)
        print(response.text)
        data = json.loads(response.text)
        print("克隆结果："+data['audio_url'])
        return data['audio_url']
    except Exception as e:
        print(f"错误：{e}")
def clean_markdown(text):
    # 去除粗体和斜体标记，例如 **text**、*text*
    text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", text)

    # 去除标题标记，例如 #、##、###、####
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)

    # 去除列表标记 - 和 *
    text = re.sub(r"^[\-\*]\s*", "", text, flags=re.MULTILINE)

    return text.strip()
def get_image_base64_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        image_data = response.content
        base64_image = base64.b64encode(image_data).decode('utf-8')
        return base64_image
    else:
        print("Error getting image")
        return JsonResponse({"result": False},status=500)
def encode_image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')


def compress_image(file, max_size_kb=1024):
    # 加载原始图像
    img = Image.open(file)
    buffer = BytesIO()

    # JPEG模式如果是 PNG 之类格式需要转换
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    # 压缩循环
    quality = 85
    while True:
        buffer.seek(0)
        buffer.truncate()
        img.save(buffer, format="JPEG", quality=quality)
        size_kb = buffer.tell() / 1024
        if size_kb <= max_size_kb or quality <= 20:
            break
        quality -= 5

    return base64.b64encode(buffer.getvalue()).decode('utf-8')
def image_model(base64_image,text):
    # ✅使用ai模型分析文字和图片
    try:
        client = OpenAI(api_key="sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
                        base_url="https://api.siliconflow.cn/v1")
        response = client.chat.completions.create(
            model="Qwen/Qwen2.5-VL-72B-Instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        },
                        {
                            "type": "text",
                            "text": f"{text}"
                        }
                    ]
                }
            ],
            temperature=0.3,
            stream=True,
        )
        return response
    except Exception as e:
        print(e)
        return JsonResponse({"result": False},status=500)
def sanitize_filename1(name):
    return "".join(re.findall(r'[\u4e00-\u9fff\w -]', name)).strip()
def event_stream(response):
    full_reply_content = ""
    first_chunk = True  # 标记是否是第一次输出

    for chunk in response:
        if chunk.choices[0].delta.content:
            part = chunk.choices[0].delta.content

            if first_chunk:
                # 第一次输出时加上开头标识
                part = "（Ai生成）" + part
                first_chunk = False

            full_reply_content += part
            print(full_reply_content)

            yield part

# def download_file(file_url, filename):
#     try:
#         response = requests.get(file_url, timeout=10)
#         if response.status_code == 200:
#             file_path = os.path.join(base_dir, filename)
#             with open(file_path, 'wb') as f:
#                 f.write(response.content)
#             print(f"✅ 下载成功：{file_path}")
#         else:
#             print(f"❌ 下载失败 {filename}，状态码：{response.status_code}")
#     except Exception as e:
#         print(f"❌ 下载异常 {filename}：{e}")
def read_docx(file_path):
    try:
        doc = Document(file_path)
        contents = "\n".join([para.text for para in doc.paragraphs])
        print(contents)
        return contents
    except Exception as e:
        print(e)
        return f"读取 Word 出错：{str(e)}"
def read_pdf(file_path):
    global text
    try:
        doc = pymupdf.open(file_path)  # open a document
        for page in doc:  # iterate the document pages
            text = page.get_text()  # get plain text (is in UTF-8)
            text += text
        print(text)
        return text
    except Exception as e:
        print(e)
        return f"读取 Pdf 出错：{str(e)}"

def read_pdf_to_markdown_with_images(file_path):
    doc = pymupdf.open(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "pdf_images")
    os.makedirs(image_dir, exist_ok=True)

    host = "http://139.196.235.10:8005"  # 你的图片基础 URL
    md_text = ""
    image_count = 0

    for page_num in range(len(doc)):
        page = doc[page_num]

        # 提取文字（按段落）
        text = page.get_text("blocks")  # [(x0, y0, x1, y1, "text", block_no, block_type)]
        text_blocks = sorted(text, key=lambda b: (b[1], b[0]))  # sort by y, then x

        for block in text_blocks:
            block_text = block[4].strip()
            if not block_text:
                continue

            # 模拟标题识别（示意：首行粗体+大号字体判断可扩展）
            if len(block_text) < 50 and block_text.endswith(":"):
                md_text += f"\n## {block_text}\n\n"
            else:
                # 每段加缩进（Markdown中无原生缩进，统一用两个空格代替）
                md_text += f"  {block_text}\n\n"

        # 提取图片
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            image_count += 1
            image_name = f"pdf_image_{page_num+1}_{image_count}.{image_ext}"
            image_path = os.path.join(image_dir, image_name)
            with open(image_path, "wb") as f:
                f.write(image_bytes)

            # 插入图片 Markdown
            md_text += f"![image]({host}/media/pdf_images/{image_name})\n\n"

    doc.close()
    return md_text

def read_docx_to_markdown_with_images_in_order(file_path):
    doc = Document(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "doc_images")
    os.makedirs(image_dir, exist_ok=True)

    host = "http://139.196.235.10:8005"
    md_text = ""
    image_count = 0

    body = doc._element.body
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    }

    for child in body.iterchildren():
        tag = etree.QName(child).localname

        if tag == "p":
            # 查找图片
            blips = child.findall(".//a:blip", namespaces=nsmap)
            if blips:
                for blip in blips:
                    rId = blip.get(f"{{{nsmap['r']}}}embed")
                    image_part = doc.part.related_parts[rId]
                    image_data = image_part.blob
                    image_name = f"image_{image_count + 1}.png"
                    image_path = os.path.join(image_dir, image_name)
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    md_text += f"![image]({host}/media/doc_images/{image_name})\n\n"
                    image_count += 1
                continue

            # 普通段落
            paragraph = None
            for p in doc.paragraphs:
                if p._p == child:
                    paragraph = p
                    break

            if paragraph:
                style = paragraph.style.name.lower()
                if "heading" in style:
                    level = style.replace("heading ", "")
                    md_text += f"\n{'#' * int(level)} {paragraph.text.strip()}\n\n"
                elif paragraph.text.strip().startswith(("-", "*", "•")):
                    md_text += f"- {paragraph.text.strip().lstrip('-•*')}\n\n"
                else:
                    # 段落文字+加粗/斜体处理 + 缩进
                    line = "  "  # 两个空格表示缩进
                    for run in paragraph.runs:
                        t = run.text.replace('\n', ' ').strip()
                        if not t:
                            continue
                        if run.bold and run.italic:
                            line += f"***{t}***"
                        elif run.bold:
                            line += f"**{t}**"
                        elif run.italic:
                            line += f"*{t}*"
                        else:
                            line += t
                    md_text += line + "\n\n"

    return md_text
def audio_to_text(tmp_audio):
    url = "https://api.siliconflow.cn/v1/audio/transcriptions"
    token = "sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru"
    headers = {
        "Authorization": f"Bearer {token}",
        # 不需要指定 Content-Type，requests 会自动添加正确的 multipart boundary
    }
    model = {
        "model": "FunAudioLLM/SenseVoiceSmall",
    }
    with open(tmp_audio.name, "rb") as f:
        files = {
            "file": (os.path.basename(tmp_audio.name), f, "audio/wav")
        }
        response = requests.post(url, headers=headers, data=model, files=files)
        result = json.loads(response.text)
    print(result['text'])
    return result['text']
def get_response(video_url):
    if "douyin" in video_url:
        type_text = "抖音"
        url = "http://api.moreapi.cn/api/douyin/aweme_detail"
    elif "kuaishou" in video_url:
        type_text = "快手"
        url = "http://api.moreapi.cn/api/ks/aweme_detail"
    elif "xhslink" in video_url:
        type_text = "小红书"
        url = "http://api.moreapi.cn/api/xhs/note_detail"
    elif "toutiao" in video_url:
        type_text = "头条"
        url = "http://api.moreapi.cn/api/toutiao/aweme_detail_v2"
    # elif "b23" in dataurl:
    #     type_text = "哔哩哔哩"
    # elif "weibo" in dataurl:
    #     type_text = "微博"
    #     url = "http://api.moreapi.cn/api/weibo/post_detail"
    # elif "xigua" in dataurl:
    #     type_text = "西瓜视频"
    #     url = "http://api.moreapi.cn/api/xigua/aweme_detail"
    print(type_text)
    payload = json.dumps({
        "aweme_id": "",
        "share_text": video_url,
        "proxy": ""
    })
    headers = {
        "Authorization": "Bearer O1Y4f9r8sbNdbSqzmpb5MUk3jMS98Hs6exTLosz8bYK0SQyyiQS6nlV2kDDVMghX",
        'Content-Type': 'application/json'
    }
    print(url)
    try:
        response = requests.request("POST", url, headers=headers, data=payload)
        text_s = json.loads(response.text)
        print(text_s)
        print(json.dumps(text_s, indent=2, ensure_ascii=False))
        return text_s,type_text
    except Exception as e:
        print(e)
        return JsonResponse({"result": False},status=500)
def text_model(text):
    client = OpenAI(api_key="sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
                    base_url="https://api.siliconflow.cn/v1")
    response = client.chat.completions.create(
        model='deepseek-ai/DeepSeek-V2.5',
        messages=[
            {'role': 'user',
             'content': f"{text}"}
        ],
        stream=True,
        temperature=0.3,
    )
    return response
@csrf_exempt
def get_comment(request):
        if request.method == 'POST':
            # ✅获取前段数据
            data = json.loads(request.body)
            print(data)
            number= data['count']
            st_url = data['url']
            print(number)
            text_type = ",".join(data["selecttext_list"])
            print(text_type)
            # ✅正则表达式提取文本中的链接
            pattern = r'https?://[^\s，。、！？“”‘’<>（）【】()"\']+'
            match = re.search(pattern, st_url)
            if match:
                video_url = match.group()
                print("提取到的视频地址:", video_url)
            else:
                print("未找到视频地址")
                return JsonResponse({"result": False},status=500)
            # ✅sh api解析视频地址
            try:
                print("执行第一方案")
                text_s, type = get_response(video_url)
                if type == '抖音':
                    url_list = text_s.get('data', {}).get('aweme_detail', {}).get('video', {}).get('play_addr', {}).get(
                        'url_list', [])
                    cover_list = text_s.get('data', {}).get('aweme_detail', {}).get('video', {}).get('cover', {}).get(
                        'url_list', [])
                    v_url = url_list[2] if len(url_list) > 2 else ""
                    cover = cover_list[1] if len(cover_list) > 1 else ""
                    print(v_url)
                    print(cover)

                    title_name = text_s.get('data', {}).get('aweme_detail', {}).get('desc', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)

                elif type == '小红书':
                    note_card = text_s.get('data', {}).get('response_body', {}).get('data', {}).get('items', [{}])[
                        0].get(
                        'note_card', {})

                    v_url = note_card.get('video', {}).get('media', {}).get('stream', {}).get('h264', [{}])[0].get(
                        'master_url', '')
                    image_list = note_card.get('image_list', [{}])[0].get('info_list', [{}])
                    cover = image_list[0].get('url', '') if image_list else ''

                    print(v_url)
                    print(cover)

                    title_name = note_card.get('title', '')
                    desc_name = note_card.get('desc', '')
                    title = sanitize_filename1(title_name)
                    desc = sanitize_filename1(desc_name)
                    print(title)
                    print(desc)

                elif type == '快手':
                    ks_data = text_s.get('data', [{}])[0]
                    v_url = ks_data.get('manifest', {}).get('adaptationSet', [{}])[0].get('representation', [{}])[
                        0].get(
                        'url', '')
                    cover = ks_data.get('coverUrls', [{}])[0].get('url', '')
                    print(v_url)
                    print(cover)

                    title_name = ks_data.get('caption', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)

                elif type == '头条':
                    tt_video = text_s.get('data', {}).get('data', {}).get('video', {})
                    url_list = tt_video.get('play_addr', {}).get('url_list', [])
                    cover_list = tt_video.get('origin_cover', {}).get('url_list', [])
                    v_url = url_list[0] if url_list else ""
                    cover = cover_list[0] if cover_list else ""
                    print(v_url)
                    print(cover)

                    title_name = text_s.get('data', {}).get('data', {}).get('title', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)
            except Exception as e:
                try:
                    print("请求失败或超时，执行第二方案:", e)
                    print("执行第二方案")
                    key = "6849cc2c1206b1978Wt433"
                    response = requests.get(f"https://api.yyy001.com/api/videoparse?url={video_url}", timeout=7)
                    if response.status_code == 200:
                        data1 = response.json()
                        v_url = data1.get('data', {}).get('url', None)
                        cover = data1['data']['cover']
                        title_name = data1['data']['title']
                        title = sanitize_filename1(title_name)
                        print("url:", v_url)
                        print("cover:", cover)
                        print("title:", title)
                    else:
                        print("接口返回非200状态码，执行第二方案")
                        return JsonResponse({"result": False},status=500)
                except Exception as e:
                    print(e)
                    return JsonResponse({"result": False},status=500)
            # ✅流式获取文字和图片
            if (v_url):
                try:
                    with tempfile.NamedTemporaryFile(suffix=".wav") as tmp_audio:
                        cmd = [
                            "ffmpeg",
                            "-y",
                            "-t", "30",
                            "-i", v_url,  # 🔁 直接从 URL 流式读
                            "-vn",  # 去掉视频
                            "-ar", "16000",  # Whisper 推荐采样率
                            "-ac", "1",  # 单声道
                            "-f", "wav",  # 输出格式
                            tmp_audio.name
                        ]
                        subprocess.run(cmd, check=True)
                        audio_text = audio_to_text(tmp_audio)
                        text = f"视频音频内容：{audio_text};视频标题：{title};请结合发送给你的视频封面图片和视频的音频内容以及视频的标题模拟真人用{number}个字左右评论这个短视频，要求评论必须要符合{text_type} 这几个类型要求，但是评论内容尽量不出现{text_type}这几个字"
                        print(text)
                        base64_image = get_image_base64_from_url(cover)
                        # ✅使用ai模型分析文字和图片
                        response_data = image_model(base64_image,text)
                        print(response_data)
                        return StreamingHttpResponse(event_stream(response_data), content_type='text/plain')
                except Exception as e:
                    print(e)
                    return JsonResponse({"result": False},status=500)
            else:
                text = f"视频标题：{title};请结合发送给你的视频封面图片和视频的音频内容以及视频的标题模拟真人用{number}个字左右评论这个短视频，要求评论必须要符合{type} 这几个类型要求，但是评论内容尽量不出现{type}这几个字"
                print(text)
                base64_image = get_image_base64_from_url(cover)
                # ✅使用ai模型分析文字和图片
                response_data = image_model(base64_image,text)
                print(response_data)
                return StreamingHttpResponse(event_stream(response_data), status=200,content_type='text/plain')
        else:
            return JsonResponse({"result": False},status=500)
@csrf_exempt
def get_listmodel(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        code = data["code"]
        print(code)
        type_list = []
        count_list = []
        with connection.cursor() as cursor:
            cursor.execute('select * from comment_type where user_code = %s', (code,))
            type_list = [
                (
                    row[0],
                    row[1],
                    bool(row[2]),
                    bool(row[3])
                )
                for row in cursor.fetchall()
                ]
        print(type_list)
        # with connection.cursor() as cursor:
        #     cursor.execute('select * from comment_count where user_code = %s', (code,))
        #     count_list = [
        #         (
        #             row[0],
        #             row[1],
        #             bool(row[2]),
        #             bool(row[3])
        #         )
        #         for row in cursor.fetchall()
        #     ]
        # print(count_list)
        with connection.cursor() as cursor:
            cursor.execute('select * from comment_type where isset=1 and user_code = %s', (code,))
            type_overlay_list = [
                (
                    row[0],
                    row[1],
                    bool(row[2]),
                    bool(row[3])
                )
                for row in cursor.fetchall()
            ]
        print(type_overlay_list)
        # with connection.cursor() as cursor:
        #     cursor.execute('select * from comment_count where isset=1 and user_code = %s', (code,))
        #     count_overlay_list = [
        #         (
        #             row[0],
        #             row[1],
        #             bool(row[2]),
        #             bool(row[3])
        #         )
        #         for row in cursor.fetchall()
        #     ]
        # print(count_overlay_list)
        return JsonResponse({"status": True,"type_list": type_list,"type_overlay_list": type_overlay_list,})
    else:
     return JsonResponse({"status": False},status=500)
@csrf_exempt
def switch_isset(request):
    if request.method == 'POST':
        print(111)
        data = json.loads(request.body)
        code = data["code"]
        print(data)
        with connection.cursor() as cursor:
            for i in range(len(data['list_type'])):
                    cursor.execute('update comment_type set isset=%s where user_code = %s and type_name = %s',(data['list_type'][i][2], code,data['list_type'][i][1]))
        return JsonResponse({"data": True})
    return JsonResponse({"data": False},status=500)
@csrf_exempt
def get_hospitallist(request):
    if request.method == 'GET':
        file_list = []
        path = '/www/wwwroot/DjangoProject4/data/hospital/hospital.xlsx'
        path_file = '/www/wwwroot/DjangoProject4/data/hospital/'
        for file in os.listdir(path_file):
            if file.endswith(".xlsx"):
                file_list.append({
                    "name":file,
                    "path":f"http://139.196.235.10:8005/media/hospital/{file}"
                })
            else:
                print("没有表格数据")
        return JsonResponse({"data": file_list})
    else:
        return JsonResponse({"data": False},status=500)
def get_filelist(request):
    if request.method == 'GET':
        folder_dict = {}
        file_path = '/www/wwwroot/DjangoProject4/data/file/'
        for folder in os.listdir(file_path):
            file_list = []
            folder_path = os.path.join(file_path, folder)
            if os.path.isdir(folder_path) and len(os.listdir(folder_path)):
                for file in os.listdir(folder_path):
                    if file.endswith(".pdf"):
                        path = f"{folder_path}/{file}"
                        name = os.path.splitext(file)[0]
                        content = read_pdf_to_markdown_with_images(path)
                        file_list.append({
                            "name":name,
                            "content":content
                        })
                    elif file.endswith(".docx"):
                        path = f"{folder_path}/{file}"
                        name = os.path.splitext(file)[0]
                        content = read_docx_to_markdown_with_images_in_order(path)
                        file_list.append({
                            "name":name,
                            "content":content
                        })
                    else:
                        print("出现不正确的文件类型")
            else:
                print("当前为空文件夹")
            folder_dict[folder] = file_list
        print(folder_dict)
        return JsonResponse({"data": folder_dict})
    else:
        return JsonResponse({"data": False},status=500)
@csrf_exempt
def get_image(request):
    global describe
    if request.method == 'POST':
        file = request.FILES['file']
        text = request.POST['text']
        type = request.POST['type']
        print("开始处理图片....")
        image_base64 = compress_image(file, max_size_kb=1024)
        print(image_base64)
        if (type == '生成菜单'):
            describe = f'根据 {text} 这几个要求和图中的食材，智能生成菜谱，输出结果只需要菜谱名和菜谱流程'
        elif (type == '查热量'):
            describe = f'根据 {text} 这几个要求和图中的食物智能识别种类，精准估算热量与营养成分，输出结果只需要输出食物种类和名称，热量还有营养成分'
        print(describe)
        response = image_model(image_base64,describe)
        return StreamingHttpResponse(event_stream(response), content_type='text/plain')
    return JsonResponse({"data": False},status=500)
@csrf_exempt
# def post_aiface(request):
#     global describe
#     if request.method == 'POST':
#         old_image = request.FILES['old_image']
#         face_image = request.FILES['face_image']
#         print("开始处理图片....")
#         old_image_base64 = compress_image(old_image, max_size_kb=1024)
#         face_image_base64 = compress_image(face_image, max_size_kb=1024)
#         task_id = str(uuid.uuid4())
#         # 先加载模型（必须调用一次 prepare）
#         prepare_res = requests.post("http://localhost:8050/prepare")
#         print("模型加载:", prepare_res.json())
#         # 发送换脸请求
#         payload = {
#             "id": task_id,
#             "inputImage": old_image_base64,
#             "targetFace": face_image_base64
#         }
#         res = requests.post("http://localhost:8050/task", json=payload)
#
#         if res.status_code == 200 and 'result' in res.json():
#             # 解码并保存换脸后的图像
#             result_base64 = res.json()['result']
#             with open("result1.jpg", "wb") as f:
#                 f.write(base64.b64decode(result_base64))
#             print("✅ 换脸完成，结果保存在 result1.jpg")
#         else:
#             print("❌ 请求失败")
#             print(res.status_code, res.text)
#
#         return JsonResponse({"data": True})
#     return JsonResponse({"data": False},status=500)
@csrf_exempt
def get_text(request):
    global describe
    if request.method == 'POST':
        data = json.loads(request.body)
        text = data['text']
        type = data['type']
        print(text)
        print(type)
        if(type == '姓名打分'):
            describe = f'根据 {text} 这个姓名，分析五行、音律、寓意等多维度，给出权威综合评分'
        elif(type == '起名'):
            describe = f'根据 {text} 这个姓氏进行智能取名，兼顾音义美与吉祥寓意'
        elif (type == '智能助手'):
            describe = f'{text}'
        response = text_model(describe)
        return StreamingHttpResponse(event_stream(response), content_type='text/plain')
    return JsonResponse({"data": False},status=500)
@csrf_exempt

@csrf_exempt
def get_code(request):
    global describe
    if request.method == 'POST':
        type_list = [
            '高情商','同意','幽默','支持','提问','感动','暖心','鼓励','加油','反对','质疑','批评',
            '惊讶','不可思议','羡慕','向往','求解答','召唤','讨论','标记','收藏','干货','有用',
            '求教程','求链接','分享经验','补充信息','热词玩梗','简短有力','神评论','表达喜爱','催更',
            '夸赞博主','价格咨询','产品细节追问','真人测评诉求','竞品对比','售后担忧','场景化需求',
            '追问原理求资料','周星驰式','梁朝伟式','预言','赞美','董宇辉式小作文','七言绝句',
            '散文诗歌','唐诗','宋词','歌词'
        ]
        data = json.loads(request.body)
        text = data.get('code')

        with connection.cursor() as cursor:
            cursor.execute('select * from user where code = %s', (text,))
            isexist = cursor.fetchall()

            if isexist:
                print("该激活码已注册")
                return JsonResponse({"data": False})  # 已存在返回 False
            else:
                print("执行注册")
                cursor.execute('insert into user(code) values (%s)', (text,))
                for item in type_list:
                    cursor.execute(
                        'insert into comment_type(type_name,isset,ischeck,user_code) values (%s,1,0,%s)',
                        (item, text)
                    )
                return JsonResponse({"data": True})  # 注册成功返回 True

    # 如果不是 POST 请求
    return JsonResponse({"data": False})

@csrf_exempt
def post_content(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        code = data['code']
        content = data['content']
        if(content):
            with connection.cursor() as cursor:
                cursor.execute('insert into content(content,user_code) values (%s,%s)',(content,code))
            with connection.cursor() as cursor:
                cursor.execute('select content from content where user_code=%s', (code,))
                select_data = cursor.fetchall()
                result = [row[0] for row in select_data]
                print(result)
            return JsonResponse({"data": result})
        else:
            with connection.cursor() as cursor:
                cursor.execute('select content from content where user_code=%s order by id desc', (code,))
                select_data = cursor.fetchall()
                result = [row[0] for row in select_data]
                print(result)
            print("content为空")
            return JsonResponse({"data": result})
    return JsonResponse({"data": False},status=500)

def image_to_base64(path):
    with open(path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')
@csrf_exempt
def post_aiface(request): # 建议项目启动时加载一次
    if request.method == 'POST':
        if request.method == 'POST':
            old_image = request.FILES['old_image']
            face_image = request.FILES['face_image']
            print(old_image)
            print(face_image)
            load_models()
            input_img = file_to_cv2_image(old_image)
            face_img = file_to_cv2_image(face_image)
            result_base64 = swap_face_from_cv2(input_img, face_img)
            if result_base64:
                return JsonResponse({'data': {"code":200,"data":result_base64}})
            else:
                return JsonResponse({'error': '换脸失败'}, status=500)

@csrf_exempt
def get_aiimage(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            text = data['text']
            print(f"收到前端请求:{text}")
            url = "https://api.siliconflow.cn/v1/images/generations"
            payload = {
                "model": "Kwai-Kolors/Kolors",
                "prompt": text,
                "image_size": "1024x1024",
                "batch_size": 1,
                "seed": 4999999999,
                "num_inference_steps": 20,
                "guidance_scale": 7.5,
            }
            headers = {
                "Authorization": "Bearer sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
                "Content-Type": "application/json"
            }
            response = requests.request("POST", url, json=payload, headers=headers)
            result = json.loads(response.text)
            image_url = result["data"][0]["url"]
            # image = get_image_base64_from_url(image_url)
            print(image_url)
            return JsonResponse({"data": image_url})
        except Exception as e:
            print(e)
            return JsonResponse({"data": False},status=500)
    else:
        return JsonResponse({"data": False},status=500)

@csrf_exempt
def get_unmarkvideo(request):
    global v_url
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            text = data['text']
            pattern = r'https?://[^\s，。、！？“”‘’<>（）【】()"\']+'
            match = re.search(pattern, text)
            if match:
                video_url = match.group()
                print("提取到的视频地址:", video_url)
            else:
                print("未找到视频地址")
                return JsonResponse({"result": False},status=500)
            try:
                print("执行第一方案")
                text_s, type = get_response(video_url)
                if type == '抖音':
                    url_list = text_s.get('data', {}).get('aweme_detail', {}).get('video', {}).get('play_addr', {}).get(
                        'url_list', [])
                    cover_list = text_s.get('data', {}).get('aweme_detail', {}).get('video', {}).get('cover', {}).get(
                        'url_list', [])
                    v_url = url_list[2] if len(url_list) > 2 else ""
                    cover = cover_list[1] if len(cover_list) > 1 else ""
                    print(v_url)
                    print(cover)

                    title_name = text_s.get('data', {}).get('aweme_detail', {}).get('desc', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)

                elif type == '小红书':
                    note_card = text_s.get('data', {}).get('response_body', {}).get('data', {}).get('items', [{}])[
                        0].get(
                        'note_card', {})

                    v_url = note_card.get('video', {}).get('media', {}).get('stream', {}).get('h264', [{}])[0].get(
                        'master_url', '')
                    image_list = note_card.get('image_list', [{}])[0].get('info_list', [{}])
                    cover = image_list[0].get('url', '') if image_list else ''

                    print(v_url)
                    print(cover)

                    title_name = note_card.get('title', '')
                    desc_name = note_card.get('desc', '')
                    title = sanitize_filename1(title_name)
                    desc = sanitize_filename1(desc_name)
                    print(title)
                    print(desc)

                elif type == '快手':
                    ks_data = text_s.get('data', [{}])[0]
                    v_url = ks_data.get('manifest', {}).get('adaptationSet', [{}])[0].get('representation', [{}])[
                        0].get(
                        'url', '')
                    cover = ks_data.get('coverUrls', [{}])[0].get('url', '')
                    print(v_url)
                    print(cover)

                    title_name = ks_data.get('caption', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)

                elif type == '头条':
                    tt_video = text_s.get('data', {}).get('data', {}).get('video', {})
                    url_list = tt_video.get('play_addr', {}).get('url_list', [])
                    cover_list = tt_video.get('origin_cover', {}).get('url_list', [])
                    v_url = url_list[0] if url_list else ""
                    cover = cover_list[0] if cover_list else ""
                    print(v_url)
                    print(cover)

                    title_name = text_s.get('data', {}).get('data', {}).get('title', '')
                    title = sanitize_filename1(title_name)
                    print(title_name)
            except Exception as e:
                try:
                    print("请求失败或超时，执行第二方案:", e)
                    print("执行第二方案")
                    key = "6849cc2c1206b1978Wt433"
                    response = requests.get(f"https://api.yyy001.com/api/videoparse?url={video_url}", timeout=7)
                    if response.status_code == 200:
                        data1 = response.json()
                        v_url = data1.get('data', {}).get('url', None)
                        cover = data1['data']['cover']
                        title_name = data1['data']['title']
                        title = sanitize_filename1(title_name)
                        print("url:", v_url)
                        print("cover:", cover)
                        print("title:", title)
                    else:
                        print("接口返回非200状态码，执行第二方案")
                        return JsonResponse({"result": False}, status=500)
                except Exception as e:
                    print(e)
                    return JsonResponse({"result": False}, status=500)
            return JsonResponse({"data": v_url})
        except Exception as e:
            print(e)
            return JsonResponse({"result": False},status=500)
    else:
        return JsonResponse({"result": False}, status=500)

@csrf_exempt
def post_audio(request):
    if request.method == 'POST':
        try:
            text= request.POST.get('text')
            file = request.FILES.get('file')
            print(f"text:{text}")
            print(f"file:{file}")
            result  = voice_copy(file,text)
            return JsonResponse({'data': result})
        except Exception as e:
            print(e)
            return JsonResponse({'error': '获取失败'}, status=500)
    else:
        return JsonResponse({'error': '获取失败'}, status=500)


@csrf_exempt
def voice_list(request):
    if request.method != 'GET':
        return JsonResponse({"status": 0, "msg": "仅支持GET请求", "data": []})

    try:
        base_path = "/www/wwwroot/DjangoProject4/data/voice"
        domain = request.build_absolute_uri('/')[:-1]  # 自动获取域名
        media_url_prefix = "/media/voice/"

        if not os.path.exists(base_path):
            return JsonResponse({"status": 0, "msg": "目录不存在", "data": []})

        audio_exts = (".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a")
        files = [f for f in os.listdir(base_path) if f.lower().endswith(audio_exts)]
        urls = [domain + media_url_prefix + f for f in files]

        return JsonResponse({"status": 1, "msg": "success", "data": urls})
    except Exception as e:
        return JsonResponse({"status": 0, "msg": str(e), "data": []})


