import json
import subprocess
import tempfile
import uuid
from io import BytesIO
import pymupdf
from docx import Document
from django.http import StreamingHttpResponse
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.db import connection
from PIL import Image
import os
import re
from openai import OpenAI
import requests
import base64
from lxml import etree
import sys


# Face swap utilities
from commentapp.magic.face import load_models, file_to_cv2_image, swap_face_from_cv2


# =========================
# Helper functions (renamed & simplified)
# =========================

def clone_voice_sync(audio, text):
    try:
        url = 'https://aivoiceclonefree.com/api/instant/clone-sync'
        api_key = 'sk-DIFcSSOXhFYwweBgPr3VEIF9CXhIslW1'
        req_data = {'text': text, 'api_key': api_key, 'type': 2}
        files = {'audio': audio}
        resp = requests.post(url, data=req_data, files=files)
        data = json.loads(resp.text)
        return data.get('audio_url')
    except Exception as exc:
        print(f"clone_voice_sync error: {exc}")
        return None


def strip_markdown(text):
    text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[\-\*]\s*", "", text, flags=re.MULTILINE)
    return text.strip()


def fetch_image_b64(url):
    resp = requests.get(url)
    if resp.status_code == 200:
        return base64.b64encode(resp.content).decode('utf-8')
    print("fetch_image_b64: error status", resp.status_code)
    return None


def encode_image_b64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')


def compress_image_to_b64(file, max_size_kb=1024):
    img = Image.open(file)
    buffer = BytesIO()
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    quality = 85
    while True:
        buffer.seek(0)
        buffer.truncate()
        img.save(buffer, format="JPEG", quality=quality)
        size_kb = buffer.tell() / 1024
        if size_kb <= max_size_kb or quality <= 20:
            break
        quality -= 5
    return base64.b64encode(buffer.getvalue()).decode('utf-8')


def gen_image_text_stream(base64_image, text):
    try:
        client = OpenAI(api_key="sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
                        base_url="https://api.siliconflow.cn/v1")
        response = client.chat.completions.create(
            model="Qwen/Qwen2.5-VL-72B-Instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        },
                        {"type": "text", "text": f"{text}"}
                    ]
                }
            ],
            temperature=0.3,
            stream=True,
        )
        return response
    except Exception as exc:
        print("gen_image_text_stream error:", exc)
        return None


def sanitize_filename(name):
    return "".join(re.findall(r'[\u4e00-\u9fff\w -]', name)).strip()


def stream_chunks(response):
    full_reply = ""
    first = True
    for chunk in response:
        if chunk.choices[0].delta.content:
            part = chunk.choices[0].delta.content
            if first:
                part = "（Ai生成）" + part
                first = False
            full_reply += part
            yield part


def read_docx_text(file_path):
    try:
        doc = Document(file_path)
        contents = "\n".join([para.text for para in doc.paragraphs])
        return contents
    except Exception as exc:
        print(exc)
        return f"读取 Word 出错：{str(exc)}"


def read_pdf_text(file_path):
    try:
        doc = pymupdf.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as exc:
        print(exc)
        return f"读取 Pdf 出错：{str(exc)}"


def pdf_to_markdown_with_images(file_path):
    doc = pymupdf.open(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "pdf_images")
    os.makedirs(image_dir, exist_ok=True)
    host = "http://139.196.235.10:8005"
    md_text = ""
    image_count = 0
    for page_num in range(len(doc)):
        page = doc[page_num]
        text_blocks = sorted(page.get_text("blocks"), key=lambda b: (b[1], b[0]))
        for block in text_blocks:
            block_text = block[4].strip()
            if not block_text:
                continue
            if len(block_text) < 50 and block_text.endswith(":"):
                md_text += f"\n## {block_text}\n\n"
            else:
                md_text += f"  {block_text}\n\n"
        images = page.get_images(full=True)
        for _, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_count += 1
            image_name = f"pdf_image_{page_num+1}_{image_count}.{image_ext}"
            image_path = os.path.join(image_dir, image_name)
            with open(image_path, "wb") as f:
                f.write(image_bytes)
            md_text += f"![image]({host}/media/pdf_images/{image_name})\n\n"
    doc.close()
    return md_text


def docx_to_markdown_with_images_in_order(file_path):
    doc = Document(file_path)
    image_dir = os.path.join(settings.MEDIA_ROOT, "doc_images")
    os.makedirs(image_dir, exist_ok=True)
    host = "http://139.196.235.10:8005"
    md_text = ""
    image_count = 0
    body = doc._element.body
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    }
    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag == "p":
            blips = child.findall(".//a:blip", namespaces=nsmap)
            if blips:
                for blip in blips:
                    rId = blip.get(f"{{{nsmap['r']}}}embed")
                    image_part = doc.part.related_parts[rId]
                    image_data = image_part.blob
                    image_name = f"image_{image_count + 1}.png"
                    image_path = os.path.join(image_dir, image_name)
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    md_text += f"![image]({host}/media/doc_images/{image_name})\n\n"
                    image_count += 1
                continue
            paragraph = None
            for p in doc.paragraphs:
                if p._p == child:
                    paragraph = p
                    break
            if paragraph:
                style = paragraph.style.name.lower()
                if "heading" in style:
                    level = style.replace("heading ", "")
                    md_text += f"\n{'#' * int(level)} {paragraph.text.strip()}\n\n"
                elif paragraph.text.strip().startswith(("-", "*", "•")):
                    md_text += f"- {paragraph.text.strip().lstrip('-•*')}\n\n"
                else:
                    line = "  "
                    for run in paragraph.runs:
                        t = run.text.replace('\n', ' ').strip()
                        if not t:
                            continue
                        if run.bold and run.italic:
                            line += f"***{t}***"
                        elif run.bold:
                            line += f"**{t}**"
                        elif run.italic:
                            line += f"*{t}*"
                        else:
                            line += t
                    md_text += line + "\n\n"
    return md_text


def transcribe_audio(tmp_audio):
    url = "https://api.siliconflow.cn/v1/audio/transcriptions"
    token = "sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru"
    headers = {"Authorization": f"Bearer {token}"}
    model = {"model": "FunAudioLLM/SenseVoiceSmall"}
    with open(tmp_audio.name, "rb") as f:
        files = {"file": (os.path.basename(tmp_audio.name), f, "audio/wav")}
        resp = requests.post(url, headers=headers, data=model, files=files)
        result = json.loads(resp.text)
    return result.get('text', '')


def fetch_video_parse(video_url):
    if "douyin" in video_url:
        platform = "抖音"
        url = "http://api.moreapi.cn/api/douyin/aweme_detail"
    elif "kuaishou" in video_url:
        platform = "快手"
        url = "http://api.moreapi.cn/api/ks/aweme_detail"
    elif "xhslink" in video_url:
        platform = "小红书"
        url = "http://api.moreapi.cn/api/xhs/note_detail"
    elif "toutiao" in video_url:
        platform = "头条"
        url = "http://api.moreapi.cn/api/toutiao/aweme_detail_v2"
    else:
        platform = ""
        url = ""
    payload = json.dumps({"aweme_id": "", "share_text": video_url, "proxy": ""})
    headers = {
        "Authorization": "Bearer O1Y4f9r8sbNdbSqzmpb5MUk3jMS98Hs6exTLosz8bYK0SQyyiQS6nlV2kDDVMghX",
        'Content-Type': 'application/json'
    }
    try:
        resp = requests.request("POST", url, headers=headers, data=payload)
        data = json.loads(resp.text)
        return data, platform
    except Exception as exc:
        print(exc)
        return None, platform


def gen_text_stream(text):
    client = OpenAI(api_key="sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
                    base_url="https://api.siliconflow.cn/v1")
    response = client.chat.completions.create(
        model='deepseek-ai/DeepSeek-V2.5',
        messages=[{'role': 'user', 'content': f"{text}"}],
        stream=True,
        temperature=0.3,
    )
    return response


# =========================
# Small utilities for complex views
# =========================

def extract_url(text):
    pattern = r'https?://[^\s，。、！？“”‘’<>（）【】()"\']+'
    match = re.search(pattern, text)
    return match.group() if match else None


def select_media_from_moreapi(data, platform):  # 从 moreapi 结构中提取媒体信息
    downloadable_url = ""  # 可下载直链
    cover_image_url = ""  # 封面图直链
    video_title = ""  # 视频标题
    if platform == '抖音':  # 抖音解析
        url_list = data.get('data', {}).get('aweme_detail', {}).get('video', {}).get('play_addr', {}).get('url_list', [])  # 视频直链列表
        cover_list = data.get('data', {}).get('aweme_detail', {}).get('video', {}).get('cover', {}).get('url_list', [])  # 封面直链列表
        downloadable_url = url_list[2] if len(url_list) > 2 else ""  # 选用较高清直链
        cover_image_url = cover_list[1] if len(cover_list) > 1 else ""  # 选用较高清封面
        title_name = data.get('data', {}).get('aweme_detail', {}).get('desc', '')  # 原始标题
        video_title = sanitize_filename(title_name)  # 标题清洗
    elif platform == '小红书':  # 小红书解析
        note_card = data.get('data', {}).get('response_body', {}).get('data', {}).get('items', [{}])[0].get('note_card', {})  # 笔记卡片
        downloadable_url = note_card.get('video', {}).get('media', {}).get('stream', {}).get('h264', [{}])[0].get('master_url', '')  # 视频直链
        image_list = note_card.get('image_list', [{}])[0].get('info_list', [{}])  # 图片列表
        cover_image_url = image_list[0].get('url', '') if image_list else ''  # 封面直链
        video_title = sanitize_filename(note_card.get('title', ''))  # 标题清洗
    elif platform == '快手':  # 快手解析
        ks_data = data.get('data', [{}])[0]  # 结构根
        downloadable_url = ks_data.get('manifest', {}).get('adaptationSet', [{}])[0].get('representation', [{}])[0].get('url', '')  # 视频直链
        cover_image_url = ks_data.get('coverUrls', [{}])[0].get('url', '')  # 封面直链
        video_title = sanitize_filename(ks_data.get('caption', ''))  # 标题清洗
    elif platform == '头条':  # 头条解析
        tt_video = data.get('data', {}).get('data', {}).get('video', {})  # 视频节点
        url_list = tt_video.get('play_addr', {}).get('url_list', [])  # 视频直链列表
        cover_list = tt_video.get('origin_cover', {}).get('url_list', [])  # 封面直链列表
        downloadable_url = url_list[0] if url_list else ""  # 选第一个直链
        cover_image_url = cover_list[0] if cover_list else ""  # 选第一个封面
        video_title = sanitize_filename(data.get('data', {}).get('data', {}).get('title', ''))  # 标题清洗
    return downloadable_url, cover_image_url, video_title  # 返回三要素


def fallback_parse_api(video_url):  # 第二方案兜底解析
    try:
        resp = requests.get(f"https://api.yyy001.com/api/videoparse?url={video_url}", timeout=7)  # 调用兜底解析API
        if resp.status_code == 200:  # 成功
            data = resp.json()  # JSON解析
            downloadable_url = data.get('data', {}).get('url', None)  # 可下载直链
            cover_image_url = data.get('data', {}).get('cover')  # 封面直链
            video_title = sanitize_filename(data.get('data', {}).get('title', ''))  # 标题清洗
            return downloadable_url, cover_image_url, video_title  # 返回三要素
        return None, None, None  # 状态码非200
    except Exception as exc:
        print("fallback_parse_api error:", exc)  # 打印错误
        return None, None, None  # 异常兜底


def download_audio_30s_ffmpeg(video_url):
    with tempfile.NamedTemporaryFile(suffix=".wav") as tmp_audio:
        cmd = [
            "ffmpeg", "-y",
            "-t", "30",
            "-i", video_url,
            "-vn", "-ar", "16000", "-ac", "1", "-f", "wav",
            tmp_audio.name
        ]
        subprocess.run(cmd, check=True)
        return tmp_audio


def build_comment_prompt(audio_text, title, number, comment_type_labels):  # 组装评论提示词
    return (  # 返回完整提示
        f"视频音频内容：{audio_text};视频标题：{title};请结合发送给你的视频封面图片和视频的音频内容以及视频的标题"
        f"模拟真人用{number}个字左右评论这个短视频，要求评论必须要符合{comment_type_labels} 这几个类型要求，"
        f"但是评论内容尽量不出现{comment_type_labels}这几个字"
    )


def image_to_base64(path):
    with open(path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')


# =========================
# Views (external names kept; internals refactored)
# =========================

@csrf_exempt
def get_comment(request):  # 生成短视频评论
    if request.method != 'POST':  # 仅支持POST
        return JsonResponse({"result": False}, status=500)  # 方法不符
    data = json.loads(request.body)  # 读取JSON体
    number = data.get('count')  # 评论字数目标
    shared_text = data.get('url', '')  # 分享文本（可能包含链接）
    comment_type_labels = ",".join(data.get("selecttext_list", []))  # 评论类型标签
    video_url = extract_url(shared_text)  # 提取视频URL
    if not video_url:  # 无URL
        return JsonResponse({"result": False}, status=500)  # 返回失败
    try:
        parse_data, platform = fetch_video_parse(video_url)  # 第一方案解析
        if parse_data:  # 解析成功
            downloadable_url, cover_image_url, video_title = select_media_from_moreapi(parse_data, platform)  # 提取三要素
        else:  # 解析失败
            downloadable_url, cover_image_url, video_title = None, None, None  # 置空
    except Exception as exc:  # 异常进入兜底
        print("fetch_video_parse error, fallback:", exc)  # 打印错误
        downloadable_url, cover_image_url, video_title = fallback_parse_api(video_url)  # 兜底解析
    if downloadable_url:  # 有可下载直链
        try:
            with tempfile.NamedTemporaryFile(suffix=".wav") as tmp_audio:  # 创建临时音频文件
                cmd = [  # ffmpeg命令
                    "ffmpeg", "-y", "-t", "30", "-i", downloadable_url,
                    "-vn", "-ar", "16000", "-ac", "1", "-f", "wav", tmp_audio.name
                ]
                subprocess.run(cmd, check=True)  # 执行抽音频
                audio_text = transcribe_audio(tmp_audio)  # 语音转文本
            prompt = build_comment_prompt(audio_text, video_title, number, comment_type_labels)  # 构造提示词
            base64_image = fetch_image_b64(cover_image_url)  # 获取封面Base64
            response_data = gen_image_text_stream(base64_image, prompt)  # 图文多模态生成
            if not response_data:  # 生成异常
                return JsonResponse({"result": False}, status=500)  # 返回失败
            return StreamingHttpResponse(stream_chunks(response_data), content_type='text/plain')  # 流式输出
        except Exception as exc:
            print(exc)  # 打印错误
            return JsonResponse({"result": False}, status=500)  # 返回失败
    else:  # 没有视频直链，仅用封面与标题生成
        prompt = (  # 构造简化提示词
            f"视频标题：{video_title};请结合发送给你的视频封面图片和视频的音频内容以及视频的标题模拟真人用{number}个字左右评论这个短视频，"
            f"要求评论必须要符合{comment_type_labels} 这几个类型要求，但是评论内容尽量不出现{comment_type_labels}这几个字"
        )
        base64_image = fetch_image_b64(cover_image_url)  # 获取封面Base64
        response_data = gen_image_text_stream(base64_image, prompt)  # 图文多模态生成
        if not response_data:  # 异常
            return JsonResponse({"result": False}, status=500)  # 返回失败
        return StreamingHttpResponse(stream_chunks(response_data), status=200, content_type='text/plain')  # 流式输出


@csrf_exempt
def get_listmodel(request):
    if request.method != 'POST':
        return JsonResponse({"status": False}, status=500)
    data = json.loads(request.body)
    code = data.get("code")
    with connection.cursor() as cursor:
        cursor.execute('select * from comment_type where user_code = %s', (code,))
        type_list = [
            (row[0], row[1], bool(row[2]), bool(row[3]))
            for row in cursor.fetchall()
        ]
    with connection.cursor() as cursor:
        cursor.execute('select * from comment_type where isset=1 and user_code = %s', (code,))
        type_overlay_list = [
            (row[0], row[1], bool(row[2]), bool(row[3]))
            for row in cursor.fetchall()
        ]
    return JsonResponse({"status": True, "type_list": type_list, "type_overlay_list": type_overlay_list})


@csrf_exempt
def switch_isset(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    code = data.get("code")
    with connection.cursor() as cursor:
        for i in range(len(data.get('list_type', []))):
            cursor.execute(
                'update comment_type set isset=%s where user_code = %s and type_name = %s',
                (data['list_type'][i][2], code, data['list_type'][i][1])
            )
    return JsonResponse({"data": True})


@csrf_exempt
def get_hospitallist(request):
    if request.method != 'GET':
        return JsonResponse({"data": False}, status=500)
    file_list = []
    path_file = '/www/wwwroot/DjangoProject4/data/hospital/'
    for file in os.listdir(path_file):
        if file.endswith(".xlsx"):
            file_list.append({"name": file, "path": f"http://139.196.235.10:8005/media/hospital/{file}"})
    return JsonResponse({"data": file_list})


def get_filelist(request):
    if request.method != 'GET':
        return JsonResponse({"data": False}, status=500)
    folder_dict = {}
    file_path = '/www/wwwroot/DjangoProject4/data/file/'
    for folder in os.listdir(file_path):
        file_list = []
        folder_path = os.path.join(file_path, folder)
        if os.path.isdir(folder_path) and len(os.listdir(folder_path)):
            for file in os.listdir(folder_path):
                if file.endswith(".pdf"):
                    path = f"{folder_path}/{file}"
                    name = os.path.splitext(file)[0]
                    content = pdf_to_markdown_with_images(path)
                    file_list.append({"name": name, "content": content})
                elif file.endswith(".docx"):
                    path = f"{folder_path}/{file}"
                    name = os.path.splitext(file)[0]
                    content = docx_to_markdown_with_images_in_order(path)
                    file_list.append({"name": name, "content": content})
        folder_dict[folder] = file_list
    return JsonResponse({"data": folder_dict})


@csrf_exempt
def get_image(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    file = request.FILES['file']
    text = request.POST['text']
    typ = request.POST['type']
    image_base64 = compress_image_to_b64(file, max_size_kb=1024)
    if typ == '生成菜单':
        prompt = f'根据 {text} 这几个要求和图中的食材，智能生成菜谱，输出结果只需要菜谱名和菜谱流程'
    elif typ == '查热量':
        prompt = f'根据 {text} 这几个要求和图中的食物智能识别种类，精准估算热量与营养成分，输出结果只需要输出食物种类和名称，热量还有营养成分'
    else:
        prompt = text
    response = gen_image_text_stream(image_base64, prompt)
    if not response:
        return JsonResponse({"data": False}, status=500)
    return StreamingHttpResponse(stream_chunks(response), content_type='text/plain')


@csrf_exempt
def get_text(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    text = data.get('text')
    typ = data.get('type')
    if typ == '姓名打分':
        prompt = f'根据 {text} 这个姓名，分析五行、音律、寓意等多维度，给出权威综合评分'
    elif typ == '起名':
        prompt = f'根据 {text} 这个姓氏进行智能取名，兼顾音义美与吉祥寓意'
    elif typ == '智能助手':
        prompt = f'{text}'
    else:
        prompt = f'{text}'
    response = gen_text_stream(prompt)
    return StreamingHttpResponse(stream_chunks(response), content_type='text/plain')


@csrf_exempt
def get_code(request):
    if request.method != 'POST':
        return JsonResponse({"data": False})
    default_types = [
        '高情商','同意','幽默','支持','提问','感动','暖心','鼓励','加油','反对','质疑','批评',
        '惊讶','不可思议','羡慕','向往','求解答','召唤','讨论','标记','收藏','干货','有用',
        '求教程','求链接','分享经验','补充信息','热词玩梗','简短有力','神评论','表达喜爱','催更',
        '夸赞博主','价格咨询','产品细节追问','真人测评诉求','竞品对比','售后担忧','场景化需求',
        '追问原理求资料','周星驰式','梁朝伟式','预言','赞美','董宇辉式小作文','七言绝句',
        '散文诗歌','唐诗','宋词','歌词'
    ]
    data = json.loads(request.body)
    code = data.get('code')
    with connection.cursor() as cursor:
        cursor.execute('select * from user where code = %s', (code,))
        isexist = cursor.fetchall()
        if isexist:
            return JsonResponse({"data": False})
        cursor.execute('insert into user(code) values (%s)', (code,))
        for item in default_types:
            cursor.execute(
                'insert into comment_type(type_name,isset,ischeck,user_code) values (%s,1,0,%s)',
                (item, code)
            )
    return JsonResponse({"data": True})


@csrf_exempt
def post_content(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    data = json.loads(request.body)
    code = data.get('code')
    content = data.get('content')
    if content:
        with connection.cursor() as cursor:
            cursor.execute('insert into content(content,user_code) values (%s,%s)', (content, code))
        with connection.cursor() as cursor:
            cursor.execute('select content from content where user_code=%s', (code,))
            result = [row[0] for row in cursor.fetchall()]
        return JsonResponse({"data": result})
    else:
        with connection.cursor() as cursor:
            cursor.execute('select content from content where user_code=%s order by id desc', (code,))
            result = [row[0] for row in cursor.fetchall()]
        return JsonResponse({"data": result})


@csrf_exempt
def post_aiface(request):  # 建议项目启动时加载一次
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    old_image = request.FILES['old_image']
    face_image = request.FILES['face_image']
    load_models()
    input_img = file_to_cv2_image(old_image)
    face_img = file_to_cv2_image(face_image)
    result_base64 = swap_face_from_cv2(input_img, face_img)
    if result_base64:
        return JsonResponse({'data': {"code": 200, "data": result_base64}})
    return JsonResponse({'error': '换脸失败'}, status=500)


@csrf_exempt
def get_aiimage(request):
    if request.method != 'POST':
        return JsonResponse({"data": False}, status=500)
    try:
        data = json.loads(request.body)
        text = data.get('text')
        url = "https://api.siliconflow.cn/v1/images/generations"
        payload = {
            "model": "Kwai-Kolors/Kolors",
            "prompt": text,
            "image_size": "1024x1024",
            "batch_size": 1,
            "seed": 4999999999,
            "num_inference_steps": 20,
            "guidance_scale": 7.5,
        }
        headers = {
            "Authorization": "Bearer sk-ebbcpdkopzbtwmeyedlqepvdeppbbkpgllhqudjuolvirxru",
            "Content-Type": "application/json"
        }
        resp = requests.request("POST", url, json=payload, headers=headers)
        result = json.loads(resp.text)
        image_url = result["data"][0]["url"]
        return JsonResponse({"data": image_url})
    except Exception as exc:
        print(exc)
        return JsonResponse({"data": False}, status=500)


@csrf_exempt
def get_unmarkvideo(request):  # 获取无水印视频直链
    if request.method != 'POST':  # 仅支持POST
        return JsonResponse({"result": False}, status=500)  # 方法不符
    try:
        data = json.loads(request.body)  # 读取JSON体
        text = data.get('text', '')  # 分享文本
        video_url = extract_url(text)  # 提取链接
        if not video_url:  # 无URL
            return JsonResponse({"result": False}, status=500)  # 返回失败
        try:  # 第一方案
            parse_data, platform = fetch_video_parse(video_url)  # 解析平台
            if not parse_data:  # 解析失败
                raise RuntimeError("primary parse failed")  # 抛出异常走兜底
            downloadable_url, cover_image_url, video_title = select_media_from_moreapi(parse_data, platform)  # 提取直链
        except Exception as exc:  # 兜底方案
            print("primary parse failed:", exc)  # 打印错误
            downloadable_url, cover_image_url, video_title = fallback_parse_api(video_url)  # 兜底解析
            if not downloadable_url:  # 仍失败
                return JsonResponse({"result": False}, status=500)  # 返回失败
        return JsonResponse({"data": downloadable_url})  # 返回直链
    except Exception as exc:  # 总异常兜底
        print(exc)  # 打印错误
        return JsonResponse({"result": False}, status=500)  # 返回失败


@csrf_exempt
def post_audio(request):
    if request.method != 'POST':
        return JsonResponse({'error': '获取失败'}, status=500)
    try:
        text = request.POST.get('text')
        file = request.FILES.get('file')
        result = clone_voice_sync(file, text)
        if not result:
            return JsonResponse({'error': '获取失败'}, status=500)
        return JsonResponse({'data': result})
    except Exception as exc:
        print(exc)
        return JsonResponse({'error': '获取失败'}, status=500)


@csrf_exempt
def voice_list(request):
    if request.method != 'GET':
        return JsonResponse({"status": 0, "msg": "仅支持GET请求", "data": []})
    try:
        base_path = "/www/wwwroot/DjangoProject4/data/voice"
        domain = request.build_absolute_uri('/')[:-1]
        media_url_prefix = "/media/voice/"
        if not os.path.exists(base_path):
            return JsonResponse({"status": 0, "msg": "目录不存在", "data": []})
        audio_exts = (".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a")
        files = [f for f in os.listdir(base_path) if f.lower().endswith(audio_exts)]
        urls = [domain + media_url_prefix + f for f in files]
        return JsonResponse({"status": 1, "msg": "success", "data": urls})
    except Exception as exc:
        return JsonResponse({"status": 0, "msg": str(exc), "data": []})


